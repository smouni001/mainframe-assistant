# GENAIV1_CLAUDE.py — Mainframe Modernization AI Assistant (PRO VERSION)
# ========================================================================

import io
import os
import zipfile
from datetime import datetime
from functools import lru_cache
from typing import Optional, List, Tuple
import pandas as pd
import streamlit as st

# ==== LLM client ====
try:
    from langchain_anthropic import ChatAnthropic
    CLAUDE_AVAILABLE = True
except ImportError:
    ChatAnthropic = None
    CLAUDE_AVAILABLE = False

# ==== Langchain utils ====
try:
    from langchain_openai import OpenAIEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document as LCDocument
    from langchain_core.prompts import PromptTemplate
    LANGCHAIN_AVAILABLE = True
except ImportError:
    OpenAIEmbeddings = None
    FAISS = None
    LANGCHAIN_AVAILABLE = False

# ===================== CONFIGURATION =====================
st.set_page_config(
    page_title="Assistant AI Mainframe Pro",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)
hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}        /* Menu Hamburger */
    footer {visibility: hidden;}           /* Footer "Made with Streamlit" */
    header {visibility: hidden;}           /* Optionnel : cache le header */
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
# ===================== SESSION STATE INITIALIZATION =====================
# Mode 1 : Analyse documentaire
if 'doc_analysis_result' not in st.session_state:
    st.session_state.doc_analysis_result = None
if 'doc_analysis_question' not in st.session_state:
    st.session_state.doc_analysis_question = None
if 'doc_analysis_files' not in st.session_state:
    st.session_state.doc_analysis_files = []

# Mode 2 : Génération JCL
if 'jcl_result' not in st.session_state:
    st.session_state.jcl_result = None
if 'jcl_language' not in st.session_state:
    st.session_state.jcl_language = None
if 'jcl_filename' not in st.session_state:
    st.session_state.jcl_filename = None

# Mode 3 : Test COBOL
if 'cobol_test_result' not in st.session_state:
    st.session_state.cobol_test_result = None
if 'cobol_test_scenarios' not in st.session_state:
    st.session_state.cobol_test_scenarios = None
if 'cobol_module_name' not in st.session_state:
    st.session_state.cobol_module_name = None
if 'cobol_test_name' not in st.session_state:
    st.session_state.cobol_test_name = None

# Mode 4 : RGC Analysis
if 'rgc_analysis_result' not in st.session_state:
    st.session_state.rgc_analysis_result = None
if 'rgc_df_config' not in st.session_state:
    st.session_state.rgc_df_config = None
if 'rgc_resume_raw' not in st.session_state:
    st.session_state.rgc_resume_raw = None
if 'rgc_stats_dict' not in st.session_state:
    st.session_state.rgc_stats_dict = None
if 'rgc_uploaded_filename' not in st.session_state:
    st.session_state.rgc_uploaded_filename = None
if 'show_advanced_analysis' not in st.session_state:
    st.session_state.show_advanced_analysis = False
if 'rgc_df_summary' not in st.session_state:
    st.session_state.rgc_df_summary = None

# Mode 5 : Pseudo-Code Conversion
if 'pseudo_result' not in st.session_state:
    st.session_state.pseudo_result = None
if 'pseudo_target_language' not in st.session_state:  # ← RENOMMÉ
    st.session_state.pseudo_target_language = None
if 'pseudo_generated_program_name' not in st.session_state:  # ← RENOMMÉ
    st.session_state.pseudo_generated_program_name = None
if 'pseudo_source_code' not in st.session_state:  # ← RENOMMÉ
    st.session_state.pseudo_source_code = None    

# Mode 6 : Extraction Règles de Gestion
if 'business_rules_result' not in st.session_state:
    st.session_state.business_rules_result = None
if 'business_rules_source_code' not in st.session_state:
    st.session_state.business_rules_source_code = None
if 'business_rules_language' not in st.session_state:
    st.session_state.business_rules_language = None
if 'business_rules_filename' not in st.session_state:
    st.session_state.business_rules_filename = None

# Mode 7 : Application Analyzer
if 'analyzer_results' not in st.session_state:
    st.session_state.analyzer_results = None
if 'analyzer_uploaded_files' not in st.session_state:
    st.session_state.analyzer_uploaded_files = []
if 'analyzer_dependency_graph' not in st.session_state:
    st.session_state.analyzer_dependency_graph = None
# ✅ APRÈS
if 'analyzer_computed_metrics' not in st.session_state:
    st.session_state.analyzer_computed_metrics = None
if 'analyzer_generated_report' not in st.session_state:
    st.session_state.analyzer_generated_report = None

# ===================== CUSTOM CSS PRO =====================
st.markdown("""
<style>
    /* ===== THEME PRINCIPAL ===== */
    :root {
        --primary-color: #4EA3FF;
        --secondary-color: #0066CC;
        --accent-color: #00D4FF;
        --bg-dark: #0E1117;
        --bg-card: #1E1E2E;
        --text-primary: #FFFFFF;
        --text-secondary: #B8B8C8;
        --success: #00D9A5;
        --warning: #FFB020;
        --error: #FF4757;
    }
    
    /* ===== BACKGROUND ===== */
    .stApp {
        background: linear-gradient(135deg, #0E1117 0%, #1a1d29 100%);
    }
    
    /* ===== HEADER PRO ===== */
    .ai-header {
        background: linear-gradient(135deg, #1E1E2E 0%, #2B2D3E 100%);
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 8px 32px rgba(78, 163, 255, 0.15);
        border: 1px solid rgba(78, 163, 255, 0.2);
        position: relative;
        overflow: hidden;
    }
    
    .ai-header::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(78, 163, 255, 0.1), transparent);
        animation: shine 3s infinite;
    }
    
    @keyframes shine {
        0% { left: -100%; }
        100% { left: 100%; }
    }
    
    .ai-header h1 {
        font-size: 2.5rem;
        font-weight: 800;
        margin: 0;
        background: linear-gradient(135deg, #4EA3FF 0%, #00D4FF 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        position: relative;
        z-index: 1;
    }
    
    .ai-header .subtitle {
        color: var(--text-secondary);
        font-size: 1rem;
        margin-top: 0.5rem;
        font-weight: 300;
        letter-spacing: 1px;
    }
    
    /* ===== CARDS GLASSMORPHISM ===== */
    .glass-card {
        background: rgba(30, 30, 46, 0.7);
        backdrop-filter: blur(10px);
        border-radius: 12px;
        padding: 1.5rem;
        border: 1px solid rgba(78, 163, 255, 0.15);
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
        margin-bottom: 1.5rem;
        transition: all 0.3s ease;
    }
    
    .glass-card:hover {
        border-color: rgba(78, 163, 255, 0.4);
        box-shadow: 0 12px 40px rgba(78, 163, 255, 0.2);
        transform: translateY(-2px);
    }
    
    /* ===== INFO BOXES ===== */
    .info-box {
        background: linear-gradient(135deg, rgba(78, 163, 255, 0.1) 0%, rgba(0, 212, 255, 0.1) 100%);
        border-left: 4px solid var(--primary-color);
        padding: 1rem 1.5rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: var(--text-primary);
        font-size: 0.95rem;
        line-height: 1.6;
    }
    
    .success-box {
        background: linear-gradient(135deg, rgba(0, 217, 165, 0.1) 0%, rgba(0, 255, 200, 0.1) 100%);
        border-left-color: var(--success);
    }
    
    .warning-box {
        background: linear-gradient(135deg, rgba(255, 176, 32, 0.1) 0%, rgba(255, 200, 100, 0.1) 100%);
        border-left-color: var(--warning);
    }
    
    .error-box {
        background: linear-gradient(135deg, rgba(255, 71, 87, 0.1) 0%, rgba(255, 100, 120, 0.1) 100%);
        border-left-color: var(--error);
    }
    
    /* ===== BUTTONS PRO ===== */
    .stButton>button {
        background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%);
        color: white;
        font-weight: 600;
        border-radius: 10px;
        border: none;
        padding: 0.75rem 2rem;
        font-size: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(78, 163, 255, 0.3);
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, var(--secondary-color) 0%, var(--primary-color) 100%);
        box-shadow: 0 6px 25px rgba(78, 163, 255, 0.5);
        transform: translateY(-2px);
    }
    
    .stButton>button:active {
        transform: translateY(0);
    }
    
    .stButton>button:disabled {
        background: linear-gradient(135deg, #2a2a3a 0%, #3a3a4a 100%);
        box-shadow: none;
        opacity: 0.5;
    }
    
    /* ===== SIDEBAR PRO ===== */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1a1d29 0%, #0E1117 100%);
        border-right: 1px solid rgba(78, 163, 255, 0.15);
    }
    
    [data-testid="stSidebar"] .stRadio > label {
        background: rgba(78, 163, 255, 0.05);
        padding: 0.5rem 1rem;
        border-radius: 8px;
        margin: 0.25rem 0;
        transition: all 0.2s ease;
    }
    
    [data-testid="stSidebar"] .stRadio > label:hover {
        background: rgba(78, 163, 255, 0.15);
    }
    
    /* ===== INPUTS ===== */
    .stTextInput input, .stTextArea textarea {
        background: rgba(30, 30, 46, 0.6) !important;
        border: 1px solid rgba(78, 163, 255, 0.3) !important;
        border-radius: 8px !important;
        color: var(--text-primary) !important;
        padding: 0.75rem !important;
        transition: all 0.3s ease !important;
    }
    
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--primary-color) !important;
        box-shadow: 0 0 0 2px rgba(78, 163, 255, 0.2) !important;
    }
    
    /* ===== FILE UPLOADER ===== */
    [data-testid="stFileUploader"] {
        background: rgba(30, 30, 46, 0.4);
        border: 2px dashed rgba(78, 163, 255, 0.3);
        border-radius: 12px;
        padding: 2rem;
        transition: all 0.3s ease;
    }
    
    [data-testid="stFileUploader"]:hover {
        border-color: var(--primary-color);
        background: rgba(78, 163, 255, 0.05);
    }
    
    /* ===== DATAFRAME ===== */
    .stDataFrame {
        border-radius: 12px;
        overflow: hidden;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
    }
    
    /* ===== CODE BLOCKS ===== */
    .stCodeBlock {
        border-radius: 12px;
        border: 1px solid rgba(78, 163, 255, 0.2);
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
    }
    
    /* ===== SPINNER ===== */
    .stSpinner > div {
        border-top-color: var(--primary-color) !important;
    }
    
    /* ===== TABS ===== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: rgba(30, 30, 46, 0.5);
        padding: 0.5rem;
        border-radius: 12px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: transparent;
        border-radius: 8px;
        color: var(--text-secondary);
        padding: 0.75rem 1.5rem;
        transition: all 0.3s ease;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%);
        color: white;
    }
    
    /* ===== FOOTER PRO ===== */
    .footer-pro {
        margin-top: 4rem;
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        background: linear-gradient(135deg, #1E1E2E 0%, #2B2D3E 100%);
        border: 1px solid rgba(78, 163, 255, 0.2);
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
    }
    
    .footer-title {
        color: var(--primary-color);
        font-weight: 700;
        font-size: 1.1rem;
        margin-bottom: 1rem;
    }
    
    .footer-team {
        color: var(--text-secondary);
        font-size: 0.9rem;
        line-height: 1.8;
    }
    
    .footer-team span {
        color: var(--accent-color);
        font-weight: 600;
    }
    
    /* ===== ANIMATIONS ===== */
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .glass-card, .info-box {
        animation: fadeIn 0.5s ease-out;
    }
    
    /* ===== RESPONSIVE ===== */
    @media (max-width: 768px) {
        .ai-header h1 {
            font-size: 1.8rem;
        }
        
        .glass-card {
            padding: 1rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# ===================== HEADER PRO =====================
st.markdown("""
<div class="ai-header">
    <h1>🤖 Assistant AI Mainframe</h1>
    <div class="subtitle">IBM A&T</div>
</div>
""", unsafe_allow_html=True)

# ===================== LANGUE =====================
lang_choice = st.sidebar.selectbox(
    "🌐 Language / Langue", 
    ["Français", "English"], 
    index=0,
    help="Sélectionnez votre langue préférée"
)
LANG_FR = (lang_choice == "Français")

def T(fr: str, en: str) -> str:
    return fr if LANG_FR else en

# ===================== API KEY =====================
ANTHROPIC_API_KEY = st.secrets["ANTHROPIC_API_KEY"]

if not CLAUDE_AVAILABLE or not ANTHROPIC_API_KEY:
    st.markdown("""
    <div class="error-box">
        ⚠️ <strong>LLM Claude indisponible</strong><br>
        Installez <code>langchain-anthropic</code> et configurez votre clé API
    </div>
    """, unsafe_allow_html=True)
    st.stop()

# ===================== PROMPT ENGINE =====================
import yaml

@lru_cache(maxsize=1)
def load_prompt_engine(path: str = "PromptEngine.yaml") -> dict:
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
        else:
            st.markdown(f"""
            <div class="warning-box">
                ℹ️ Fichier <code>{path}</code> introuvable. Utilisation des prompts par défaut.
            </div>
            """, unsafe_allow_html=True)
            return {}
    except Exception as e:
        st.markdown(f"""
        <div class="error-box">
            ❌ Erreur lors du chargement de <code>{path}</code>: {e}
        </div>
        """, unsafe_allow_html=True)
        return {}

def get_prompt(section: str, lang_key: Optional[str] = None, **kwargs) -> str:
    pe = load_prompt_engine()
    defaults = pe.get("defaults", {})
    invalid_message = defaults.get("invalid_message", "❌ Code invalide")
    no_jcl = defaults.get("no_jcl_in_test_mode", "Ne pas générer de JCL")

    if section == "ANALYSE_DOC":
        node = pe.get("ANALYSE_DOC", {}).get("GENERIQUE", "")
    elif section == "PSEUDOCODE_CONVERSION":
        node = pe.get("PSEUDOCODE_CONVERSION", {}).get(lang_key.upper() if lang_key else "COBOL", "")
    elif lang_key:
        node = pe.get(section, {}).get(lang_key.upper(), "")
    else:
        node = ""

    if not node:
        if section == "ANALYSE_DOC":
            node = "Rôle: Analyste documentaire.\n{document_content}"
        elif section == "PSEUDOCODE_CONVERSION":
            node = """Tu es un Expert Mainframe Senior.
Convertis ce pseudo-code en {lang_key} :
{pseudocode}
Programme: {program_name}"""
        elif section == "JCL_GENERATION":
            node = "Rôle: Expert JCL.\nCode:\n{source_code}"
        else:
            node = "Rôle: QA Test.\nModule:{module_base}\nCode:\n{source_code}"

    # AJOUT DES PARAMÈTRES PAR DÉFAUT
    default_params = {
        'invalid_message': invalid_message,
        'no_jcl_in_test_mode': no_jcl,
        'lang_key': lang_key if lang_key else 'COBOL',
        'lang': lang_key if lang_key else 'COBOL'
    }
    
    all_params = {**default_params, **kwargs}

    try:
        return node.format(**all_params)
    except KeyError as e:
        st.error(f"❌ Paramètre manquant dans le prompt : {e}")
        st.info(f"Paramètres disponibles : {list(all_params.keys())}")
        st.code(node, language="yaml")
        return node

# ===================== HELPERS =====================
def llm_client(max_tokens: int = 4000, temperature: float = 0.2):
    if not CLAUDE_AVAILABLE or not ANTHROPIC_API_KEY:
        return None
    try:
        return ChatAnthropic(
            api_key=ANTHROPIC_API_KEY,
            temperature=temperature,
            max_tokens=max_tokens,
            model_name="claude-3-haiku-20240307"
        )
    except Exception as e:
        st.error(f"❌ Erreur client Claude: {e}")
        return None

def _normalize_text(s: str) -> str:
    return " ".join(s.replace("\r", "\n").split())

def read_pdf_bytes(data: bytes) -> str:
    try:
        import PyPDF2
        text = []
        with io.BytesIO(data) as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text.append(extracted)
        result = _normalize_text("\n".join(text))
        if result.strip():
            return result
    except:
        pass
    
    try:
        import pdfplumber
        text = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text.append(extracted)
        return _normalize_text("\n".join(text))
    except:
        return ""

def read_docx_bytes(data: bytes) -> str:
    try:
        import docx
        with io.BytesIO(data) as fh:
            doc = docx.Document(fh)
            return _normalize_text("\n".join(p.text for p in doc.paragraphs))
    except:
        return ""

def read_txt_bytes(data: bytes, encoding: str = "utf-8") -> str:
    try:
        return _normalize_text(data.decode(encoding, errors="ignore"))
    except:
        return ""

def extract_document_content(uploaded_file) -> str:
    if not uploaded_file:
        return ""
    
    name = (uploaded_file.name or "").lower()
    
    try:
        raw = uploaded_file.read()
    except Exception as e:
        st.error(f"❌ Erreur de lecture: {e}")
        return ""

    def safe_truncate(text: str, max_chars: int = 150_000) -> str:
        if len(text) > max_chars:
            st.markdown(f"""
            <div class="warning-box">
                ⚠️ Contenu tronqué à {max_chars:,} caractères
            </div>
            """, unsafe_allow_html=True)
        return text[:max_chars]

    if name.endswith(".pdf"):
        return safe_truncate(read_pdf_bytes(raw))
    if name.endswith(".docx"):
        return safe_truncate(read_docx_bytes(raw))
    if name.endswith(".txt"):
        return safe_truncate(read_txt_bytes(raw))

    if name.endswith(".zip"):
        texts = []
        try:
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                for info in z.infolist():
                    n = info.filename.lower()
                    if n.endswith((".pdf", ".docx", ".txt")) and not info.is_dir():
                        try:
                            data = z.read(info)
                            if n.endswith(".pdf"):
                                texts.append(read_pdf_bytes(data))
                            elif n.endswith(".docx"):
                                texts.append(read_docx_bytes(data))
                            elif n.endswith(".txt"):
                                texts.append(read_txt_bytes(data))
                            
                            if sum(len(t) for t in texts) > 200_000:
                                break
                        except Exception as e:
                            st.warning(f"⚠️ Impossible de lire {info.filename}")
                            continue
        except zipfile.BadZipFile:
            st.markdown("""
            <div class="error-box">
                ❌ Fichier ZIP invalide ou corrompu
            </div>
            """, unsafe_allow_html=True)
            return ""
        
        return safe_truncate("\n\n---\n\n".join(t for t in texts if t.strip()))

    return ""

def detect_critical_conflicts(df: pd.DataFrame) -> pd.DataFrame:
    """
    Post-traitement pour garantir la détection des conflits critiques
    même si le LLM les a manqués.
    """
    if df.empty or 'Nom Programme' not in df.columns:
        return df
    
    required_cols = ['Nom Programme', 'Environnement', 'Couloir', 'Niveau de Risque']
    if not all(col in df.columns for col in required_cols):
        return df
    
    df['_env_couloir_key'] = df['Environnement'].astype(str) + '_' + df['Couloir'].astype(str)
    
    for prog_name, group in df.groupby('Nom Programme'):
        if len(group) > 1:
            duplicates = group[group.duplicated(subset=['_env_couloir_key'], keep=False)]
            
            if not duplicates.empty:
                indices = duplicates.index
                for idx in indices:
                    df.at[idx, 'Niveau de Risque'] = '🚨 Conflit de duplication dans le même environnement'
            else:
                unique_envs = group['Environnement'].nunique()
                unique_couloirs = group['Couloir'].nunique()
                
                if unique_envs > 1:
                    for idx in group.index:
                        current_risk = df.at[idx, 'Niveau de Risque']
                        if '🚨' not in str(current_risk):
                            df.at[idx, 'Niveau de Risque'] = '⚠️ Programme dupliqué (multi-environnements)'
                elif unique_couloirs > 1:
                    for idx in group.index:
                        current_risk = df.at[idx, 'Niveau de Risque']
                        if '🚨' not in str(current_risk):
                            df.at[idx, 'Niveau de Risque'] = '⚠️ Risque de modification (multi-couloir)'
    
    df.drop(columns=['_env_couloir_key'], inplace=True)
    
    return df
# ===================== APPLICATION ANALYZER - CLASSES & PARSERS =====================
import re
import json
from dataclasses import dataclass, asdict
from typing import Dict, Set
from collections import defaultdict

@dataclass
class ProgramInfo:
    name: str
    path: str
    lines: int
    complexity: int
    reads: List[str]
    writes: List[str]
    calls: List[str]
    paragraphs: int
    io_operations: int

@dataclass
class JCLStepInfo:
    name: str
    job: str
    program: str
    conds: List[str]
    dd_names: List[str]
    datasets: List[str]

@dataclass
class DatasetInfo:
    name: str
    type: str
    produced_by: List[str]
    consumed_by: List[str]

class COBOLParser:
    """Parser simple pour COBOL"""
    
    @staticmethod
    def parse(content: str, filename: str) -> ProgramInfo:
        lines = content.split('\n')
        
        # Comptage basique
        total_lines = len([l for l in lines if l.strip() and not l.strip().startswith('*')])
        
        # Extraction du nom de programme
        prog_name = filename.replace('.cbl', '').replace('.cob', '').upper()
        for line in lines:
            if 'PROGRAM-ID' in line.upper():
                match = re.search(r'PROGRAM-ID\.\s+(\w+)', line.upper())
                if match:
                    prog_name = match.group(1)
                    break
        
        # Comptage des paragraphes
        paragraphs = len(re.findall(r'^\s*[A-Z0-9\-]+\s*\.\s*$', content, re.MULTILINE))
        
        # Opérations I/O
        io_ops = (
            content.upper().count('READ ') + 
            content.upper().count('WRITE ') + 
            content.upper().count('REWRITE ') +
            content.upper().count('DELETE ')
        )
        
        # Fichiers en lecture
        reads = []
        for line in lines:
            if 'SELECT' in line.upper() and 'ASSIGN' in line.upper():
                match = re.search(r'SELECT\s+(\w+)', line.upper())
                if match:
                    reads.append(match.group(1))
        
        # Fichiers en écriture (même logique)
        writes = list(set(re.findall(r'WRITE\s+(\w+)', content.upper())))
        
        # Appels de programmes
        calls = list(set(re.findall(r'CALL\s+[\'"](\w+)[\'"]', content.upper())))
        
        # Complexité cyclomatique approximée
        complexity = (
            content.upper().count('IF ') +
            content.upper().count('PERFORM ') +
            content.upper().count('GO TO ') +
            content.upper().count('EVALUATE ')
        )
        
        return ProgramInfo(
            name=prog_name,
            path=filename,
            lines=total_lines,
            complexity=max(1, complexity),
            reads=reads,
            writes=writes,
            calls=calls,
            paragraphs=paragraphs,
            io_operations=io_ops
        )

class JCLParser:
    """Parser simple pour JCL"""
    
    @staticmethod
    def parse(content: str, filename: str) -> List[JCLStepInfo]:
        steps = []
        lines = content.split('\n')
        
        current_job = "UNKNOWN"
        current_step = None
        current_pgm = None
        current_dds = []
        current_conds = []
        current_datasets = []
        
        for line in lines:
            line = line.strip()
            
            # Job card
            if line.startswith('//') and 'JOB' in line:
                match = re.match(r'//(\w+)\s+JOB', line)
                if match:
                    current_job = match.group(1)
            
            # Step card
            elif line.startswith('//') and 'EXEC' in line:
                # Sauvegarder le step précédent
                if current_step:
                    steps.append(JCLStepInfo(
                        name=current_step,
                        job=current_job,
                        program=current_pgm or "UNKNOWN",
                        conds=current_conds,
                        dd_names=current_dds,
                        datasets=current_datasets
                    ))
                
                # Nouveau step
                match = re.match(r'//(\w+)\s+EXEC', line)
                if match:
                    current_step = match.group(1)
                    current_dds = []
                    current_conds = []
                    current_datasets = []
                    
                    # Programme
                    pgm_match = re.search(r'PGM=(\w+)', line)
                    if pgm_match:
                        current_pgm = pgm_match.group(1)
                    
                    # Condition
                    cond_match = re.search(r'COND=\(([^)]+)\)', line)
                    if cond_match:
                        current_conds.append(cond_match.group(1))
            
            # DD statement
            elif line.startswith('//') and 'DD' in line and current_step:
                match = re.match(r'//(\w+)\s+DD', line)
                if match:
                    dd_name = match.group(1)
                    current_dds.append(dd_name)
                    
                    # Dataset name
                    dsn_match = re.search(r'DSN=([^,\s]+)', line)
                    if dsn_match:
                        current_datasets.append(dsn_match.group(1))
        
        # Dernier step
        if current_step:
            steps.append(JCLStepInfo(
                name=current_step,
                job=current_job,
                program=current_pgm or "UNKNOWN",
                conds=current_conds,
                dd_names=current_dds,
                datasets=current_datasets
            ))
        
        return steps

class ApplicationAnalyzer:
    """Analyseur principal"""
    
    def __init__(self):
        self.programs: Dict[str, ProgramInfo] = {}
        self.jcl_steps: List[JCLStepInfo] = []
        self.datasets: Dict[str, DatasetInfo] = {}
        self.dependencies: Dict[str, Set[str]] = defaultdict(set)
    
    def analyze_files(self, files: List[Tuple[str, bytes]]):
        """Analyse tous les fichiers uploadés"""
        
        for filename, content_bytes in files:
            try:
                content = content_bytes.decode('utf-8', errors='ignore')
                
                if filename.endswith(('.cbl', '.cob')):
                    prog_info = COBOLParser.parse(content, filename)
                    self.programs[prog_info.name] = prog_info
                
                elif filename.endswith(('.jcl', '.txt')):
                    steps = JCLParser.parse(content, filename)
                    self.jcl_steps.extend(steps)
                    
                    # Construire dépendances
                    for i, step in enumerate(steps):
                        if i > 0:
                            prev_step = steps[i-1]
                            self.dependencies[step.name].add(prev_step.name)
            
            except Exception as e:
                st.warning(f"⚠️ Erreur parsing {filename}: {e}")
                continue
        
        # Construire datasets
        self._build_datasets()
    
    def _build_datasets(self):
        """Construit le registre des datasets"""
        for step in self.jcl_steps:
            for dsn in step.datasets:
                if dsn not in self.datasets:
                    self.datasets[dsn] = DatasetInfo(
                        name=dsn,
                        type="SEQUENTIAL",
                        produced_by=[],
                        consumed_by=[]
                    )
                
                # Heuristique : si DD OUTPUT/REPORT → produced, sinon consumed
                if any(dd in ['OUTPUT', 'REPORT', 'SORTOUT'] for dd in step.dd_names):
                    self.datasets[dsn].produced_by.append(step.name)
                else:
                    self.datasets[dsn].consumed_by.append(step.name)
    
    def compute_metrics(self) -> Dict:
        """Calcule les métriques globales"""
        total_lines = sum(p.lines for p in self.programs.values())
        avg_complexity = sum(p.complexity for p in self.programs.values()) / max(1, len(self.programs))
        
        return {
            "total_programs": len(self.programs),
            "total_jcl_steps": len(self.jcl_steps),
            "total_datasets": len(self.datasets),
            "total_lines": total_lines,
            "avg_complexity": round(avg_complexity, 2),
            "max_complexity": max((p.complexity for p in self.programs.values()), default=0),
            "total_io_operations": sum(p.io_operations for p in self.programs.values())
        }
    
    def build_dependency_graph(self):
        """Construit le graphe NetworkX"""
        try:
            import networkx as nx
        except ImportError:
            st.error("❌ Installez networkx : pip install networkx")
            return None
        
        G = nx.DiGraph()
        
        # Ajouter les noeuds
        for prog_name in self.programs.keys():
            G.add_node(prog_name, type='PROGRAM')
        
        for step in self.jcl_steps:
            G.add_node(step.name, type='STEP')
        
        # Ajouter les arêtes
        for step in self.jcl_steps:
            if step.program in self.programs:
                G.add_edge(step.name, step.program, type='EXECUTES')
        
        for target, sources in self.dependencies.items():
            for source in sources:
                G.add_edge(source, target, type='PRECEDES')
        
        return G
    
    def generate_report_json(self) -> Dict:
        """Génère le rapport JSON complet"""
        graph_dict = {
            "nodes": list(self.programs.keys()) + [s.name for s in self.jcl_steps],
            "edges": [[s, t] for s, targets in self.dependencies.items() for t in targets]
        }
        
        return {
            "meta": {
                "project": "MainframeAnalysis",
                "analyzed_at": datetime.now().isoformat()
            },
            "programs": [asdict(p) for p in self.programs.values()],
            "jcl_steps": [asdict(s) for s in self.jcl_steps],
            "datasets": [asdict(d) for d in self.datasets.values()],
            "dependency_graph": graph_dict,
            "metrics_summary": self.compute_metrics()
        }
# ===================== REVERSE ENGINEERING CLASSES =====================
import re
import json
from dataclasses import dataclass, asdict, field
from typing import Dict, Set, List, Optional
from collections import defaultdict
from enum import Enum

class ComponentType(Enum):
    """Types de composants Mainframe"""
    COBOL_BATCH = "COBOL_BATCH"
    COBOL_CICS = "COBOL_CICS"
    COBOL_IMS = "COBOL_IMS"
    COBOL_DB2 = "COBOL_DB2"
    COBOL_HYBRID = "COBOL_HYBRID"
    JCL = "JCL"
    COPYBOOK = "COPYBOOK"
    BMS_MAP = "BMS_MAP"
    DB2_TABLE = "DB2_TABLE"
    IMS_SEGMENT = "IMS_SEGMENT"
    MQ_QUEUE = "MQ_QUEUE"
    VSAM_FILE = "VSAM_FILE"
    TRANSACTION = "TRANSACTION"

@dataclass
class CallInfo:
    """Information sur un appel"""
    caller: str
    callee: str
    call_type: str
    line_number: Optional[int] = None
    context: Optional[str] = None

@dataclass
class ProgramMetadata:
    """Métadonnées complètes d'un programme"""
    name: str
    path: str
    component_type: ComponentType
    lines: int
    paragraphs: int
    sections: int
    
    calls_static: List[str] = field(default_factory=list)
    calls_dynamic: List[str] = field(default_factory=list)
    calls_cics: List[Dict[str, str]] = field(default_factory=list)
    calls_ims: List[Dict[str, str]] = field(default_factory=list)
    calls_db2: List[str] = field(default_factory=list)
    calls_mq: List[str] = field(default_factory=list)
    
    copybooks: List[str] = field(default_factory=list)
    files_vsam: List[str] = field(default_factory=list)
    files_sequential: List[str] = field(default_factory=list)
    db2_tables: List[str] = field(default_factory=list)
    ims_segments: List[str] = field(default_factory=list)
    mq_queues: List[str] = field(default_factory=list)
    
    transactions: List[str] = field(default_factory=list)
    bms_maps: List[str] = field(default_factory=list)
    
    called_by: List[str] = field(default_factory=list)
    calls_to: List[str] = field(default_factory=list)
    
    is_orphan: bool = False
    is_critical: bool = False
    complexity_score: int = 0
    risk_level: str = "LOW"

class AdvancedCOBOLParser:
    """Parser COBOL avancé avec détection CICS/IMS/DB2/MQ"""
    
    @staticmethod
    def parse(content: str, filename: str) -> ProgramMetadata:
        lines = content.split('\n')
        
        prog_name = filename.replace('.cbl', '').replace('.cob', '').upper()
        for line in lines:
            if 'PROGRAM-ID' in line.upper():
                match = re.search(r'PROGRAM-ID\.\s+(\w+)', line.upper())
                if match:
                    prog_name = match.group(1)
                    break
        
        content_upper = content.upper()
        component_type = ComponentType.COBOL_BATCH
        
        if 'EXEC CICS' in content_upper:
            component_type = ComponentType.COBOL_CICS
        if 'DLITCBL' in content_upper or 'ENTRY \'DLITCBL\'' in content_upper:
            component_type = ComponentType.COBOL_IMS
        if 'EXEC SQL' in content_upper:
            if component_type == ComponentType.COBOL_CICS:
                component_type = ComponentType.COBOL_HYBRID
            else:
                component_type = ComponentType.COBOL_DB2
        
        total_lines = len([l for l in lines if l.strip() and not l.strip().startswith('*')])
        paragraphs = len(re.findall(r'^\s*[A-Z0-9\-]+\s*\.\s*$', content, re.MULTILINE))
        sections = content_upper.count(' SECTION.')
        
        copybooks = list(set(re.findall(r'COPY\s+(\w+)', content_upper)))
        calls_static = list(set(re.findall(r'CALL\s+[\'"](\w+)[\'"]', content_upper)))
        calls_dynamic = list(set(re.findall(r'CALL\s+(\w+)\s+USING', content_upper)))
        calls_dynamic = [c for c in calls_dynamic if c not in calls_static and c not in ['CBLTDLI', 'MQPUT', 'MQGET']]
        
        calls_cics = []
        for match in re.finditer(r'EXEC\s+CICS\s+(LINK|XCTL|START)\s+.*?PROGRAM\s*\(\s*[\'"]?(\w+)[\'"]?\s*\)', content_upper):
            calls_cics.append({'type': match.group(1), 'program': match.group(2)})
        
        transactions = list(set(re.findall(r'EXEC\s+CICS\s+START\s+TRANSID\s*\(\s*[\'"]?(\w+)[\'"]?\s*\)', content_upper)))
        bms_maps = list(set(re.findall(r'EXEC\s+CICS\s+(?:SEND|RECEIVE)\s+MAP\s*\(\s*[\'"]?(\w+)[\'"]?\s*\)', content_upper)))
        
        calls_ims = []
        ims_segments = []
        for match in re.finditer(r'CALL\s+[\'"]CBLTDLI[\'"].*?USING\s+(\w+)', content_upper):
            calls_ims.append({'type': 'DL/I', 'function': match.group(1)})
        
        ims_segments = list(set(re.findall(r'(\w+SEGMENT|\w+SEG)', content_upper)))
        
        db2_tables = []
        for match in re.finditer(r'EXEC\s+SQL\s+(?:SELECT|UPDATE|DELETE|INSERT).*?(?:FROM|INTO|UPDATE)\s+(\w+)', content_upper, re.DOTALL):
            table = match.group(1).split()[0]
            if table not in ['DUAL', 'SYSIBM']:
                db2_tables.append(table)
        db2_tables = list(set(db2_tables))
        
        mq_queues = []
        calls_mq = []
        if 'MQPUT' in content_upper or 'MQGET' in content_upper:
            calls_mq = ['MQPUT', 'MQGET']
            mq_queues = list(set(re.findall(r'QUEUE\s*\(\s*[\'"]?(\w+)[\'"]?\s*\)', content_upper)))
        
        files_vsam = []
        files_sequential = []
        
        for match in re.finditer(r'SELECT\s+(\w+)\s+ASSIGN\s+TO\s+(\w+)', content_upper):
            file_name = match.group(1)
            if 'VSAM' in content_upper or 'KSDS' in content_upper:
                files_vsam.append(file_name)
            else:
                files_sequential.append(file_name)
        
        complexity_score = (
            len(calls_static) * 2 +
            len(calls_dynamic) * 3 +
            len(calls_cics) * 2 +
            len(calls_ims) * 3 +
            len(db2_tables) * 2 +
            content_upper.count('IF ') +
            content_upper.count('PERFORM ') +
            content_upper.count('EVALUATE ')
        )
        
        risk_level = "LOW"
        if complexity_score > 50:
            risk_level = "HIGH"
        elif complexity_score > 25:
            risk_level = "MEDIUM"
        
        return ProgramMetadata(
            name=prog_name,
            path=filename,
            component_type=component_type,
            lines=total_lines,
            paragraphs=paragraphs,
            sections=sections,
            calls_static=calls_static,
            calls_dynamic=calls_dynamic,
            calls_cics=calls_cics,
            calls_ims=calls_ims,
            calls_db2=db2_tables,
            calls_mq=calls_mq,
            copybooks=copybooks,
            files_vsam=files_vsam,
            files_sequential=files_sequential,
            db2_tables=db2_tables,
            ims_segments=ims_segments,
            mq_queues=mq_queues,
            transactions=transactions,
            bms_maps=bms_maps,
            complexity_score=complexity_score,
            risk_level=risk_level
        )

class AdvancedJCLParser:
    """Parser JCL avancé"""
    
    @staticmethod
    def parse(content: str, filename: str) -> Dict:
        lines = content.split('\n')
        
        job_name = "UNKNOWN"
        steps = []
        current_step = None
        
        for line in lines:
            line = line.strip()
            
            if line.startswith('//') and 'JOB' in line:
                match = re.match(r'//(\w+)\s+JOB', line)
                if match:
                    job_name = match.group(1)
            
            elif line.startswith('//') and 'EXEC' in line:
                if current_step:
                    steps.append(current_step)
                
                match = re.match(r'//(\w+)\s+EXEC', line)
                if match:
                    step_name = match.group(1)
                    
                    pgm_match = re.search(r'PGM=(\w+)', line)
                    proc_match = re.search(r'PROC=(\w+)', line)
                    
                    current_step = {
                        'name': step_name,
                        'program': pgm_match.group(1) if pgm_match else None,
                        'proc': proc_match.group(1) if proc_match else None,
                        'cond': re.search(r'COND=\(([^)]+)\)', line).group(1) if 'COND=' in line else None,
                        'datasets': []
                    }
            
            elif line.startswith('//') and 'DD' in line and current_step:
                dsn_match = re.search(r'DSN=([^,\s]+)', line)
                if dsn_match:
                    current_step['datasets'].append(dsn_match.group(1))
        
        if current_step:
            steps.append(current_step)
        
        return {
            'job_name': job_name,
            'filename': filename,
            'steps': steps
        }

class MainframeReverseEngineer:
    """Analyseur principal de paysage applicatif"""
    
    def __init__(self):
        self.programs: Dict[str, ProgramMetadata] = {}
        self.jcl_jobs: List[Dict] = []
        self.copybooks: Set[str] = set()
        self.transactions: Dict[str, str] = {}
        self.bms_maps: Dict[str, List[str]] = defaultdict(list)
        self.db2_tables: Set[str] = set()
        self.ims_segments: Set[str] = set()
        self.mq_queues: Set[str] = set()
        self.call_graph: Dict[str, Set[str]] = defaultdict(set)
        self.reverse_call_graph: Dict[str, Set[str]] = defaultdict(set)
    
    def analyze_sources(self, files: List[Tuple[str, bytes]]):
        """Analyse tous les fichiers sources"""
        
        for filename, content_bytes in files:
            try:
                content = content_bytes.decode('utf-8', errors='ignore')
                basename = filename.split('/')[-1].lower()
                
                if basename.endswith(('.cbl', '.cob')):
                    prog_meta = AdvancedCOBOLParser.parse(content, filename)
                    self.programs[prog_meta.name] = prog_meta
                    
                    self.copybooks.update(prog_meta.copybooks)
                    self.db2_tables.update(prog_meta.db2_tables)
                    self.ims_segments.update(prog_meta.ims_segments)
                    self.mq_queues.update(prog_meta.mq_queues)
                    
                    for trans in prog_meta.transactions:
                        self.transactions[trans] = prog_meta.name
                    
                    for bms_map in prog_meta.bms_maps:
                        self.bms_maps[bms_map].append(prog_meta.name)
                
                elif basename.endswith('.jcl'):
                    jcl_data = AdvancedJCLParser.parse(content, filename)
                    self.jcl_jobs.append(jcl_data)
                
                elif basename.endswith(('.cpy', '.copy')):
                    copy_name = basename.replace('.cpy', '').replace('.copy', '').upper()
                    self.copybooks.add(copy_name)
            
            except Exception as e:
                st.warning(f"⚠️ Erreur parsing {filename}: {e}")
                continue
        
        self._build_call_graphs()
        self._detect_orphans()
        self._identify_critical_programs()
    
    def _build_call_graphs(self):
        """Construit le graphe d'appels complet"""
        
        for prog_name, prog_meta in self.programs.items():
            for callee in prog_meta.calls_static:
                self.call_graph[prog_name].add(callee)
                self.reverse_call_graph[callee].add(prog_name)
                prog_meta.calls_to.append(callee)
            
            for cics_call in prog_meta.calls_cics:
                callee = cics_call['program']
                self.call_graph[prog_name].add(callee)
                self.reverse_call_graph[callee].add(prog_name)
                prog_meta.calls_to.append(callee)
            
            for jcl in self.jcl_jobs:
                for step in jcl['steps']:
                    if step['program'] == prog_name:
                        self.reverse_call_graph[prog_name].add(f"JCL:{jcl['job_name']}")
        
        for prog_name, prog_meta in self.programs.items():
            prog_meta.called_by = list(self.reverse_call_graph.get(prog_name, []))
    
    def _detect_orphans(self):
        """Détecte les programmes non utilisés"""
        
        for prog_name, prog_meta in self.programs.items():
            if not prog_meta.called_by and prog_name not in [step['program'] for jcl in self.jcl_jobs for step in jcl['steps']]:
                prog_meta.is_orphan = True
    
    def _identify_critical_programs(self):
        """Identifie les programmes critiques (centraux)"""
        
        for prog_name, prog_meta in self.programs.items():
            if (len(prog_meta.called_by) >= 3 or 
                len(prog_meta.calls_to) >= 5 or 
                prog_meta.complexity_score > 40):
                prog_meta.is_critical = True
    
    def compute_statistics(self) -> Dict:
        """Calcule les statistiques globales"""
        
        stats = {
            'total_programs': len(self.programs),
            'cobol_batch': sum(1 for p in self.programs.values() if p.component_type == ComponentType.COBOL_BATCH),
            'cobol_cics': sum(1 for p in self.programs.values() if p.component_type == ComponentType.COBOL_CICS),
            'cobol_ims': sum(1 for p in self.programs.values() if p.component_type == ComponentType.COBOL_IMS),
            'cobol_db2': sum(1 for p in self.programs.values() if p.component_type == ComponentType.COBOL_DB2),
            'cobol_hybrid': sum(1 for p in self.programs.values() if p.component_type == ComponentType.COBOL_HYBRID),
            'total_jcl': len(self.jcl_jobs),
            'total_copybooks': len(self.copybooks),
            'total_transactions': len(self.transactions),
            'total_bms_maps': len(self.bms_maps),
            'total_db2_tables': len(self.db2_tables),
            'total_ims_segments': len(self.ims_segments),
            'total_mq_queues': len(self.mq_queues),
            'orphan_programs': sum(1 for p in self.programs.values() if p.is_orphan),
            'critical_programs': sum(1 for p in self.programs.values() if p.is_critical),
            'high_risk_programs': sum(1 for p in self.programs.values() if p.risk_level == 'HIGH'),
            'total_lines': sum(p.lines for p in self.programs.values()),
            'avg_complexity': sum(p.complexity_score for p in self.programs.values()) / max(1, len(self.programs))
        }
        
        return stats
    
    def build_dependency_graph(self):
        """Construit le graphe NetworkX complet"""
        
        try:
            import networkx as nx
        except ImportError:
            return None
        
        G = nx.DiGraph()
        
        for prog_name, prog_meta in self.programs.items():
            G.add_node(
                prog_name, 
                type='PROGRAM',
                component_type=prog_meta.component_type.value,
                is_critical=prog_meta.is_critical,
                is_orphan=prog_meta.is_orphan,
                risk_level=prog_meta.risk_level
            )
        
        for table in self.db2_tables:
            G.add_node(table, type='DB2_TABLE')
        
        for segment in self.ims_segments:
            G.add_node(segment, type='IMS_SEGMENT')
        
        for queue in self.mq_queues:
            G.add_node(queue, type='MQ_QUEUE')
        
        for caller, callees in self.call_graph.items():
            for callee in callees:
                G.add_edge(caller, callee, type='CALL')
        
        for prog_name, prog_meta in self.programs.items():
            for table in prog_meta.db2_tables:
                G.add_edge(prog_name, table, type='DB2_ACCESS')
            
            for segment in prog_meta.ims_segments:
                G.add_edge(prog_name, segment, type='IMS_ACCESS')
            
            for queue in prog_meta.mq_queues:
                G.add_edge(prog_name, queue, type='MQ_ACCESS')
        
        return G
    
    def generate_report_json(self) -> Dict:
        """Génère le rapport JSON complet"""
        
        return {
            'meta': {
                'project': 'Mainframe Reverse Engineering',
                'analyzed_at': datetime.now().isoformat(),
                'analyzer_version': '2.0'
            },
            'statistics': self.compute_statistics(),
            'programs': {
                name: {
                    **asdict(meta),
                    'component_type': meta.component_type.value
                } 
                for name, meta in self.programs.items()
            },
            'jcl_jobs': self.jcl_jobs,
            'global_resources': {
                'copybooks': list(self.copybooks),
                'transactions': self.transactions,
                'bms_maps': dict(self.bms_maps),
                'db2_tables': list(self.db2_tables),
                'ims_segments': list(self.ims_segments),
                'mq_queues': list(self.mq_queues)
            },
            'call_graph': {
                caller: list(callees) 
                for caller, callees in self.call_graph.items()
            }
        }
    
    def search_impact(self, program_name: str) -> Dict:
        """Recherche d'impact : qui sera affecté si on modifie ce programme"""
        
        if program_name not in self.programs:
            return {'error': f'Programme {program_name} non trouvé'}
        
        impacted = set()
        direct_callers = self.reverse_call_graph.get(program_name, set())
        impacted.update(direct_callers)
        
        def find_callers_recursive(prog, visited=None):
            if visited is None:
                visited = set()
            if prog in visited:
                return
            visited.add(prog)
            
            for caller in self.reverse_call_graph.get(prog, []):
                impacted.add(caller)
                find_callers_recursive(caller, visited)
        
        find_callers_recursive(program_name)
        
        return {
            'modified_program': program_name,
            'direct_impact': list(direct_callers),
            'total_impacted': list(impacted),
            'impact_count': len(impacted),
            'risk_assessment': 'HIGH' if len(impacted) > 10 else 'MEDIUM' if len(impacted) > 3 else 'LOW'
        }

# ===================== UI LABELS =====================
# ===================== UI LABELS =====================
TEXTS = {
    "Français": {
        "choose_mode": "⚙️ Mode de traitement",
        "modes": [
            "🔄 Pseudo-Code → Language Mainframe",
            "🔧 Génération JCL", 
            "🧪 Test COBOL",
            "📄 Analyse documentaire", 
            "⚙️ Analyse RGC",
            "📊 Extraction Règles de Gestion",
            "🔍 Application Analyzer"  # ← NOUVEAU MODE
        ],
    },
    "English": {
        "choose_mode": "⚙️ Processing Mode",
        "modes": [
            "🔄 Pseudo-Code → Mainframe Language",
            "🔧 JCL Generation", 
            "🧪 COBOL Testing",
            "📄 Document Analysis", 
            "⚙️ RGC Analysis",
            "📊 Business Rules Extraction",
            "🔍 Application Analyzer"  # ← NOUVEAU MODE
        ],
    }
}
TXT = TEXTS["Français"] if LANG_FR else TEXTS["English"]

# ===================== MODE SELECTOR =====================
st.sidebar.markdown("---")
mode = st.sidebar.radio(
    TXT["choose_mode"], 
    TXT["modes"], 
    index=0,
    help=T("Sélectionnez le mode de traitement souhaité", "Select your processing mode")
)

# ===================== MODE 1 : ANALYSE DOC =====================
if mode == TXT["modes"][3]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header(T("📄 Analyse Documentaire Intelligente", "📄 Intelligent Document Analysis"))
    st.markdown('</div>', unsafe_allow_html=True)

    # Upload section
    col1, col2 = st.columns(2)
    with col1:
        uploaded_zip = st.file_uploader(
            "📦 " + T("Fichier ZIP", "ZIP File"),
            type=["zip"],
            help=T("Charger un fichier ZIP contenant des documents", "Upload a ZIP file with documents"),
            key="doc_zip_uploader"
        )
    with col2:
        uploaded_pdfs = st.file_uploader(
            "📄 " + T("Fichiers PDF", "PDF Files"),
            type=["pdf"],
            accept_multiple_files=True,
            help=T("Charger un ou plusieurs PDFs", "Upload one or more PDFs"),
            key="doc_pdf_uploader"
        )

    pdf_buffers: List[Tuple[str, bytes]] = []
    
    if uploaded_zip or uploaded_pdfs:
        with st.spinner(T("📚 Traitement des documents...", "📚 Processing documents...")):
            if uploaded_zip:
                try:
                    with zipfile.ZipFile(uploaded_zip, "r") as zip_ref:
                        for name in zip_ref.namelist():
                            if name.lower().endswith(".pdf") and not name.startswith("__MACOSX"):
                                try:
                                    with zip_ref.open(name) as pdf_file:
                                        pdf_buffers.append((name, pdf_file.read()))
                                except Exception as e:
                                    st.warning(f"⚠️ {name}: {e}")
                except zipfile.BadZipFile:
                    st.markdown('<div class="error-box">❌ ZIP invalide</div>', unsafe_allow_html=True)
            
            if uploaded_pdfs:
                for f in uploaded_pdfs:
                    try:
                        pdf_buffers.append((f.name, f.read()))
                    except Exception as e:
                        st.warning(f"⚠️ {f.name}: {e}")
    
    if pdf_buffers:
        st.markdown(f"""
        <div class="success-box">
            ✅ <strong>{len(pdf_buffers)} document(s) chargé(s)</strong>
        </div>
        """, unsafe_allow_html=True)
        st.session_state.doc_analysis_files = [name for name, _ in pdf_buffers]

    # Analysis context
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    context_mode = st.radio(
        "🎯 " + T("Contexte d'analyse", "Analysis Context"),
        [T("Analyse Générique", "Generic Analysis"),
         T("Analyse CV / Profil", "CV / Profile Analysis")],
        horizontal=True,
        key="doc_context_mode"
    )

    question = st.text_area(
        "💬 " + T("Votre question", "Your question"),
        placeholder=T(
            "Ex: Quelles sont les compétences principales ?",
            "Ex: What are the main skills?"
        ),
        height=100,
        key="doc_question_input",
        value=st.session_state.doc_analysis_question if st.session_state.doc_analysis_question else ""
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # Bouton d'analyse
    analyze_button = st.button(
        "🚀 " + T("ANALYSER", "ANALYZE"), 
        disabled=not question or not pdf_buffers,
        use_container_width=True,
        key="doc_analyze_btn"
    )

    if analyze_button:
        st.session_state.doc_analysis_question = question
        
        if not pdf_buffers:
            st.markdown('<div class="error-box">❌ Aucun document chargé</div>', unsafe_allow_html=True)
        else:
            texts = []
            for name, raw in pdf_buffers:
                extracted = read_pdf_bytes(raw)
                if extracted.strip():
                    texts.append(f"=== {name} ===\n{extracted}")
            
            document_content = "\n\n---\n\n".join(texts)[:150_000]

            if not document_content.strip():
                st.markdown('<div class="error-box">❌ Aucun texte extrait</div>', unsafe_allow_html=True)
            else:
                if "CV" in context_mode or "Profile" in context_mode:
                    prompt_text = f"""
Tu es une IA experte en analyse de CV et matching de profils.

**Mission :**
- Analyse approfondie des documents fournis
- Identification des compétences (explicites et implicites)
- Évaluation complète avant réponse
- Aucun détail important omis

**Format de réponse :**
1. Réponse directe
2. Détails et justifications
3. Score (si applicable) : X/100
4. Recommandations

**Question :** {question}

**Documents :**
{document_content}
"""
                else:
                    prompt_text = f"""
Rôle : Vous êtes un Analyste Stratégique Expert, spécialisé dans l'interprétation, l'analyse et la projection d'informations basées sur des documents.

Objectif Principal : Répondre aux questions en utilisant le contenu des documents fournis, mais en allant au-delà de la simple extraction. Vous devez analyser, synthétiser et générer des réponses déductives pour éclairer l'utilisateur sur des conséquences ou des scénarios non explicitement mentionnés.

Instructions de Traitement et d'Intelligence
Analyse Complète et Déduction : Lisez tous les documents. Vous devez non seulement extraire les faits, mais aussi déduire, interpréter et extrapoler des conclusions logiques et plausibles basées sur ces faits.

Génération de Contenu : Vous êtes autorisé à inventer des réponses (dans le sens de créer des analyses, des rapports, des projections ou des scénarios) à condition qu'elles soient directement et logiquement justifiées par les données présentes dans les documents fournis.

Synthèse Stratégique : Si la question le demande (ex: "Quel impact cela pourrait-il avoir ?", "Quelle est la tendance ?"), combinez les informations de manière créative pour générer une réponse stratégique.

Capacités Avancées : Utilisez des calculs (projections, tendances, pourcentages) et des comparaisons avancées pour étayer vos déductions.

Contraintes de Réponse
Lien avec le Contexte : Toute déduction ou invention doit rester étroitement liée au contexte des documents. Si vous créez une nouvelle information, elle doit être une conséquence logique et clairement étayée par le texte source.

Absence d'Information Justificative : Si la question nécessite une déduction ou une invention qui ne peut pas être logiquement justifiée par les documents, vous devez répondre strictement par la phrase suivante, et rien d'autre :

"Information ou déduction non justifiée par les documents analysés."

Format d'Analyse : Structurez vos réponses de manière claire, concise et professionnelle :

Utilisez des listes numérotées ou à puces pour les extractions de faits.

Séparez clairement les Faits Extraits des Déductions/Scénarios Proposés.

Mettez en gras les faits critiques et les conclusions déductives.

Question : {question}

**Documents :**
{document_content}
"""

                client = llm_client(max_tokens=3000, temperature=0.2)
                if not client:
                    st.markdown('<div class="error-box">❌ Client LLM indisponible</div>', unsafe_allow_html=True)
                else:
                    with st.spinner(T("🧠 Analyse en cours...", "🧠 Analyzing...")):
                        try:
                            response = client.invoke(prompt_text)
                            result = response.content if hasattr(response, 'content') else str(response)
                            st.session_state.doc_analysis_result = result
                        except Exception as e:
                            st.error(f"❌ Erreur LLM: {e}")
                            st.session_state.doc_analysis_result = None

    # Affichage des résultats (persistant)
    if st.session_state.doc_analysis_result:
        result = st.session_state.doc_analysis_result
        
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("🧠 " + T("Réponse de l'IA", "AI Answer"))
        st.markdown(result)
        st.markdown('</div>', unsafe_allow_html=True)

        # Export buttons (toujours visibles)
        col1, col2 = st.columns(2)
        
        with col1:
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(T("Rapport d'Analyse IA", "AI Analysis Report"), 0)
                doc.add_heading(T("Question", "Question"), 1)
                doc.add_paragraph(st.session_state.doc_analysis_question)
                doc.add_heading(T("Réponse", "Answer"), 1)
                doc.add_paragraph(result)
                
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                
                st.download_button(
                    "📥 " + T("Télécharger Rapport (Word)", "Download Report (Word)"),
                    data=buf,
                    file_name=T("rapport_analyse.docx", "analysis_report.docx"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="doc_download_word"
                )
            except ImportError:
                pass
        
        with col2:
            st.download_button(
                "📥 " + T("Télécharger Rapport (TXT)", "Download Report (TXT)"),
                data=result.encode("utf-8"),
                file_name=T("rapport_analyse.txt", "analysis_report.txt"),
                use_container_width=True,
                key="doc_download_txt"
            )

# ===================== MODE 2 : JCL GENERATION =====================
elif mode == TXT["modes"][1]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header("🔧 " + T("Générateur JCL Automatique", "Automatic JCL Generator"))
    st.markdown('</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "📂 " + T("Fichier programme", "Program file"),
        type=["cbl", "cob", "pli", "pl1", "asm"],
        help=T("Sélectionnez votre fichier COBOL/PL/I/ASM", "Select your COBOL/PL/I/ASM file"),
        key="jcl_file_uploader"
    )
    
    if uploaded_file:
        st.markdown(f"""
        <div class="success-box">
            ✅ Fichier chargé : <strong>{uploaded_file.name}</strong>
        </div>
        """, unsafe_allow_html=True)
        st.session_state.jcl_filename = uploaded_file.name

    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("⚙️ " + T("Paramètres z/OS", "z/OS Parameters"))
    
    col1, col2, col3 = st.columns(3)
    with col1:
        job_name = st.text_input("🧾 JOB", value="GENJOB1", key="jcl_job_name")
        program_name = st.text_input("🏷️ " + T("Programme", "Program"), value="PROGTEST", key="jcl_program_name")
    with col2:
        pds_source = st.text_input("📁 PDS Source", value="MYUSER.COBOL.SOURCE", key="jcl_pds_source")
        loadlib = st.text_input("💾 LOADLIB", value="MYUSER.COBOL.LOAD", key="jcl_loadlib")
    with col3:
        sysout_class = st.text_input("🖨️ SYSOUT", value="A", key="jcl_sysout")
        region_size = st.text_input("💽 REGION", value="4096K", key="jcl_region")
    
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown(f"""
    <div class="info-box">
        📊 <strong>Configuration :</strong> JOB={job_name} | PGM={program_name} | PDS={pds_source}
    </div>
    """, unsafe_allow_html=True)

    generate_button = st.button(
        "🚀 " + T("GÉNÉRER JCL", "GENERATE JCL"),
        disabled=not uploaded_file,
        use_container_width=True,
        key="jcl_generate_btn"
    )

    if generate_button and uploaded_file:
        file_name = uploaded_file.name
        ext = file_name.split(".")[-1].lower()
        
        try:
            source_code = uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception as e:
            st.error(f"❌ Erreur: {e}")
            st.stop()

        lang_map = {"cbl": "COBOL", "cob": "COBOL", "pli": "PL/I", "pl1": "PL/I", "asm": "ASM"}
        lang = lang_map.get(ext)

        if not lang:
            st.markdown('<div class="error-box">⚠️ Type non supporté</div>', unsafe_allow_html=True)
        else:
            st.session_state.jcl_language = lang
            
            st.markdown(f"""
            <div class="info-box">
                🔍 Type détecté : <strong>{lang}</strong>
            </div>
            """, unsafe_allow_html=True)
            
            prompt_text = get_prompt(
                "JCL_GENERATION", lang,
                source_code=source_code,
                job_name=job_name,
                program_name=program_name,
                pds_source=pds_source,
                loadlib=loadlib,
                sysout_class=sysout_class,
                region_size=region_size
            )
            
            client = llm_client(max_tokens=2500, temperature=0.3)
            if client:
                with st.spinner(T("🧠 Génération du JCL...", "🧠 Generating JCL...")):
                    try:
                        response = client.invoke(prompt_text)
                        result = response.content if hasattr(response, 'content') else str(response)
                        st.session_state.jcl_result = result
                    except Exception as e:
                        st.error(f"❌ Erreur: {e}")
                        st.session_state.jcl_result = None

    # Affichage des résultats (persistant)
    if st.session_state.jcl_result:
        result = st.session_state.jcl_result
        lang = st.session_state.jcl_language
        
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("📘 " + T("JCL Généré", "Generated JCL"))
        st.code(result, language="jcl")
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.download_button(
            "💾 " + T("Télécharger JCL", "Download JCL"),
            data=result.encode("utf-8"),
            file_name=f"job_{lang.lower()}.jcl",
            use_container_width=True,
            key="jcl_download_btn"
        )

# ===================== MODE 3 : COBOL TEST =====================
elif mode == TXT["modes"][2]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header("🧪 " + T("Générateur de Tests COBOL", "COBOL Test Generator"))
    st.markdown('</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "📂 " + T("Module source", "Source module"),
        type=["cbl", "cob", "pli", "pl1", "asm"],
        help=T("Sélectionnez votre module", "Select your module"),
        key="cobol_file_uploader"
    )
    
    if uploaded_file:
        st.markdown(f"""
        <div class="success-box">
            ✅ Module chargé : <strong>{uploaded_file.name}</strong>
        </div>
        """, unsafe_allow_html=True)

        file_name = uploaded_file.name
        module_base = file_name.rsplit('.', 1)[0]
        ext = file_name.split(".")[-1].lower()
        
        st.session_state.cobol_module_name = module_base
        
        try:
            source_code = uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception as e:
            st.error(f"❌ Erreur: {e}")
            st.stop()

        lang_map = {"cbl": "COBOL", "cob": "COBOL", "pli": "PL/I", "pl1": "PL/I", "asm": "ASM"}
        language_type = lang_map.get(ext)

        if not language_type:
            st.markdown('<div class="error-box">⚠️ Type non supporté</div>', unsafe_allow_html=True)
            st.stop()

        module_starts_ok = module_base.upper().startswith("M")
        if not module_starts_ok:
            st.markdown(f"""
            <div class="warning-box">
                ℹ️ Le nom du module doit commencer par <strong>M</strong><br>
                Exemple : <code>MTRAITSIM.cbl</code>
            </div>
            """, unsafe_allow_html=True)

        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        default_test_name = (module_base.upper() + "TEST")
        if not default_test_name.startswith("M"):
            default_test_name = "M" + default_test_name
        default_test_name = "".join(ch for ch in default_test_name if ch.isalnum())[:8]

        test_prog_name_input = st.text_input(
            "🏷️ " + T("Nom du programme de test", "Test program name"),
            value=default_test_name,
            help=T("Commence par 'M', max 8 caractères", "Starts with 'M', max 8 chars"),
            key="cobol_test_name_input"
        )

        # Sécuriser la conversion en majuscules
        test_prog_name = test_prog_name_input.upper() if test_prog_name_input else default_test_name
        st.session_state.cobol_test_name = test_prog_name
        
        st.markdown('</div>', unsafe_allow_html=True)

        def valid_prog_name(name: str) -> Tuple[bool, str]:
            if not name:
                return False, T("Nom vide", "Empty name")
            if len(name) > 8:
                return False, T("Plus de 8 caractères", "More than 8 chars")
            if name[0] != "B":
                return False, T("Doit commencer par 'B'", "Must start with 'B'")
            if not all(ch.isalnum() for ch in name):
                return False, T("Caractères invalides", "Invalid characters")
            return True, ""

        ok_name, err_name = valid_prog_name(test_prog_name)
        if not ok_name:
            st.markdown(f"""
            <div class="warning-box">
                ℹ️ Nom invalide : {err_name}
            </div>
            """, unsafe_allow_html=True)

        st.markdown(f"""
        <div class="info-box">
            🔍 Type : <strong>{language_type}</strong> | Programme test : <strong>{test_prog_name}</strong>
        </div>
        """, unsafe_allow_html=True)

        gen_disabled = not module_starts_ok or not ok_name
        
        generate_button = st.button(
            "🚀 " + T("GÉNÉRER TEST + SCÉNARIOS", "GENERATE TEST + SCENARIOS"),
            disabled=gen_disabled,
            use_container_width=True,
            key="cobol_generate_btn"
        )
        
        if generate_button:
            prompt_text = get_prompt(
                "TEST_FACTORY", language_type,
                source_code=source_code,
                module_base=module_base,
                test_prog_name=test_prog_name
            )
            
            client = llm_client(max_tokens=4000, temperature=0.2)
            if client:
                with st.spinner(T("🧠 Génération en cours...", "🧠 Generating...")):
                    try:
                        response = client.invoke(prompt_text)
                        result = response.content if hasattr(response, 'content') else str(response)
                        
                        # Parse sections
                        cobol_test, cases_raw = "", ""
                        if "=== COBOL_TEST ===" in result:
                            after = result.split("=== COBOL_TEST ===", 1)[1]
                            if "=== CAS_DE_TEST ===" in after:
                                cobol_test, cases_raw = after.split("=== CAS_DE_TEST ===", 1)
                            else:
                                cobol_test = after

                        cobol_test_clean = cobol_test.strip()
                        
                        # Header injection
                        AUTHOR_NAME = "AYMANE"
                        DATE_STR = datetime.now().strftime("%Y/%m/%d")
                        header = f"""      *> PROGRAM-ID : {test_prog_name}
      *> AUTHOR     : {AUTHOR_NAME}
      *> DATE       : {DATE_STR}
"""
                        cobol_with_header = header + cobol_test_clean if cobol_test_clean else ""

                        # Parse scenarios
                        rows, headers = [], []
                        for ln in cases_raw.splitlines():
                            line = ln.strip()
                            if not line or line.startswith("#"):
                                continue
                            parts = [c.strip() for c in line.split("|")]
                            if not headers and "Nom_Scenario" in line:
                                headers = parts
                                continue
                            if headers and len(parts) == len(headers):
                                rows.append(parts)

                        if headers and rows:
                            df_cases = pd.DataFrame(rows, columns=headers)
                        else:
                            df_cases = pd.DataFrame(columns=[
                                "Nom_Scenario", "Condition_Entree", "Donnees_Exemple",
                                "Resultat_Attendu", "RC_Attendu", "Artefacts_Impactes", "Commentaires"
                            ])

                        # Stocker dans session_state
                        st.session_state.cobol_test_result = cobol_with_header
                        st.session_state.cobol_test_scenarios = df_cases
                        
                    except Exception as e:
                        st.error(f"❌ Erreur: {e}")
                        st.session_state.cobol_test_result = None
                        st.session_state.cobol_test_scenarios = None

    # Affichage des résultats (persistant)
    if st.session_state.cobol_test_result:
        cobol_with_header = st.session_state.cobol_test_result
        df_cases = st.session_state.cobol_test_scenarios
        test_prog_name = st.session_state.cobol_test_name
        module_base = st.session_state.cobol_module_name
        
        # Display COBOL
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("📘 " + T("Programme COBOL de test", "COBOL Test Program"))
        st.code(cobol_with_header, language="cobol")
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.download_button(
            "💾 " + T("Télécharger COBOL (.cbl)", "Download COBOL (.cbl)"),
            data=cobol_with_header.encode("utf-8"),
            file_name=f"{test_prog_name}.cbl",
            use_container_width=True,
            key="cobol_download_code"
        )

        # Display matrix
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("🧾 " + T("Matrice de scénarios", "Scenario Matrix"))
        
        if not df_cases.empty:
            st.dataframe(df_cases, use_container_width=True, height=400)
        else:
            st.warning(T("Aucun scénario détecté", "No scenarios detected"))
        
        st.markdown('</div>', unsafe_allow_html=True)

        # Export Excel/CSV
        scenarios_name = f"SCENARIOS_{module_base}"
        
        try:
            excel_buf = io.BytesIO()
            with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
                df_cases.to_excel(writer, index=False, sheet_name="SCENARIOS")
            excel_buf.seek(0)
            
            st.download_button(
                "📥 " + T("Télécharger Scénarios (Excel)", "Download Scenarios (Excel)"),
                data=excel_buf,
                file_name=f"{scenarios_name}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="cobol_download_excel"
            )
        except ImportError:
            csv_buf = io.StringIO()
            df_cases.to_csv(csv_buf, index=False, encoding="utf-8")
            csv_buf.seek(0)
            
            st.download_button(
                "📥 " + T("Télécharger Scénarios (CSV)", "Download Scenarios (CSV)"),
                data=csv_buf.getvalue().encode("utf-8"),
                file_name=f"{scenarios_name}.csv",
                use_container_width=True,
                key="cobol_download_csv"
            )

# ===================== MODE 4 : RGC ANALYSIS =====================
elif mode == TXT["modes"][4]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header("⚙️ " + T("Analyse RGC - Gestion de Configuration Logicielle", "RGC Analysis - Software Configuration Management"))
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="info-box">
        📊 <strong>Mode RGC Expert - SCM/RCM</strong><br>
        Analysez vos fichiers de configuration de programmes mainframe avec détection automatique des risques :<br>
        • Programmes dupliqués multi-environnements<br>
        • Conflits de couloir<br>
        • Alertes critiques de duplication<br>
        <br>
        <strong>Génération automatique :</strong> Tableau Excel + Résumé des risques + Métriques
    </div>
    """, unsafe_allow_html=True)

    # Upload section
    uploaded_config = st.file_uploader(
        "📄 " + T("Fichier de configuration programmes (.txt)", "Program configuration file (.txt)"),
        type=["txt"],
        help=T(
            "Format attendu : Nom_Programme | Type | Date_Création | Date_Modif | Environnement | Couloir | Modifié_Par",
            "Expected format: Program_Name | Type | Creation_Date | Modif_Date | Environment | Couloir | Modified_By"
        ),
        key="rgc_file_uploader"
    )
     # Structure ZIP recommandée
    with st.expander("📚 Structure fichier txt à importer "):
        col_left= st.columns(1)[0]
        
        with col_left:
            st.markdown("""
            **Structure txt :**
            ```
            GCL.txt
            ├── programme : 8 caractéres | Version : 3 caractéres | Type de programme : 3 caractéres | Date de création : 10 caractéres | Date de modification : 10 caractéres | Environnement : 4 Caractéres | Couloir : 2 caractéres | Modifié Par : X Caractéres 
           
            Exemple : 
            PGMCOB01V01COB2024-01-152024-03-02PRODC1ASMUNI
            
            """)
        

    if uploaded_config:
        # Vérifier si c'est un nouveau fichier
        if st.session_state.rgc_uploaded_filename != uploaded_config.name:
            st.session_state.rgc_uploaded_filename = uploaded_config.name
            # Réinitialiser les résultats pour un nouveau fichier
            st.session_state.rgc_analysis_result = None
            st.session_state.rgc_df_config = None
            st.session_state.rgc_resume_raw = None
            st.session_state.rgc_stats_dict = None
            st.session_state.show_advanced_analysis = False
            st.session_state.rgc_df_summary = None
        
        st.markdown(f"""
        <div class="success-box">
            ✅ Fichier chargé : <strong>{uploaded_config.name}</strong>
        </div>
        """, unsafe_allow_html=True)

        # Options d'analyse
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("🎯 " + T("Options d'analyse", "Analysis Options"))
        
        col1, col2, col3 = st.columns(3)
        with col1:
            detect_duplicates = st.checkbox(
                "🔍 " + T("Détection doublons", "Duplicate detection"),
                value=True,
                help=T("Identifier les programmes présents dans plusieurs environnements", 
                       "Identify programs in multiple environments"),
                key="rgc_detect_dup"
            )
        with col2:
            detect_conflicts = st.checkbox(
                "🚨 " + T("Alertes critiques", "Critical alerts"),
                value=True,
                help=T("Détecter les conflits dans le même env/couloir", 
                       "Detect conflicts in same env/couloir"),
                key="rgc_detect_conflicts"
            )
        with col3:
            generate_metrics = st.checkbox(
                "📊 " + T("Métriques détaillées", "Detailed metrics"),
                value=True,
                help=T("Générer distributions et analyses", 
                       "Generate distributions and analysis"),
                key="rgc_gen_metrics"
            )
        
        st.markdown('</div>', unsafe_allow_html=True)

        # Bouton d'analyse
        analyze_button = st.button(
            "🚀 " + T("LANCER ANALYSE RGC", "START RGC ANALYSIS"),
            use_container_width=True,
            key="rgc_analyze_btn"
        )

        # Lancer l'analyse si le bouton est cliqué
        if analyze_button:
            try:
                uploaded_config.seek(0)
                config_content = uploaded_config.read().decode("utf-8", errors="ignore")
            except Exception as e:
                st.error(f"❌ Erreur de lecture: {e}")
                st.stop()

            if not config_content.strip():
                st.markdown('<div class="error-box">❌ Fichier vide</div>', unsafe_allow_html=True)
                st.stop()

            # Préparer le prompt
            prompt_text = get_prompt(
                "RGC_ANALYSIS",
                "SCM_CONFIG",
                config_content=config_content[:150_000]
            )

            client = llm_client(max_tokens=4000, temperature=0.1)
            if not client:
                st.markdown('<div class="error-box">❌ Client LLM indisponible</div>', unsafe_allow_html=True)
                st.stop()

            with st.spinner(T("🧠 Analyse RGC en cours...", "🧠 RGC Analysis in progress...")):
                try:
                    response = client.invoke(prompt_text)
                    result = response.content if hasattr(response, 'content') else str(response)
                    st.session_state.rgc_analysis_result = result
                except Exception as e:
                    st.error(f"❌ Erreur LLM: {e}")
                    st.session_state.rgc_analysis_result = None

        # Affichage des résultats (persistant)
        if st.session_state.rgc_analysis_result:
            # Parsing des résultats (une seule fois si pas déjà fait)
            if st.session_state.rgc_df_config is None:
                result = st.session_state.rgc_analysis_result
                tableau_raw, resume_raw, stats_raw = "", "", ""
                
                if "=== TABLEAU_CONFIG ===" in result:
                    parts = result.split("=== TABLEAU_CONFIG ===", 1)[1]
                    if "=== RESUME_RISQUES ===" in parts:
                        tableau_raw, rest = parts.split("=== RESUME_RISQUES ===", 1)
                        if "=== STATISTIQUES ===" in rest:
                            resume_raw, stats_raw = rest.split("=== STATISTIQUES ===", 1)
                        else:
                            resume_raw = rest
                    else:
                        tableau_raw = parts

                # Construction du dataframe
                rows, headers = [], []

                for line in tableau_raw.splitlines():
                    line = line.strip()
                    if not line or line.startswith("-") or line.startswith("="):
                        continue
                    
                    parts = [c.strip() for c in line.split("|")]
                    
                    if len(parts) == 1 and not parts[0]:
                        continue
                    
                    if not headers and len(parts) >= 8:
                        headers = parts
                        continue
                    
                    if headers and len(parts) == len(headers):
                        rows.append(parts)

                if headers and rows:
                    if headers[0] != "N°":
                        headers.insert(0, "N°")
                        for idx, row in enumerate(rows, start=1):
                            row.insert(0, str(idx))
                    
                    df_config = pd.DataFrame(rows, columns=headers)
                    df_config = detect_critical_conflicts(df_config)
                else:
                    df_config = pd.DataFrame(columns=[
                        "N°", "Nom Programme", "Type de Programme", "Date Création", 
                        "Date Modification", "Environnement", "Couloir", 
                        "Modifié Par", "Niveau de Risque"
                    ])

                # Parser les stats
                stats_dict = {}
                for line in stats_raw.splitlines():
                    if ":" in line:
                        key, val = line.split(":", 1)
                        stats_dict[key.strip()] = val.strip()

                # Stocker dans session_state
                st.session_state.rgc_df_config = df_config
                st.session_state.rgc_resume_raw = resume_raw
                st.session_state.rgc_stats_dict = stats_dict

            # Récupérer depuis session_state
            df_config = st.session_state.rgc_df_config
            resume_raw = st.session_state.rgc_resume_raw
            stats_dict = st.session_state.rgc_stats_dict

            # Affichage tableau principal
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("📊 " + T("Tableau de Configuration", "Configuration Table"))
            
            if not df_config.empty:
                def color_risk(val):
                    if pd.isna(val) or val == "—":
                        return ''
                    elif "🚨" in str(val):
                        return 'background-color: #FF4757; color: white; font-weight: bold;'
                    elif "⚠️" in str(val):
                        return 'background-color: #FFB020; color: black;'
                    return ''
                
                styled_df = df_config.style.applymap(
                    color_risk,
                    subset=['Niveau de Risque'] if 'Niveau de Risque' in df_config.columns else []
                )
                
                st.dataframe(styled_df, use_container_width=True, height=500)
                
                # Compteurs
                if 'Niveau de Risque' in df_config.columns:
                    total = len(df_config)
                    critiques = df_config['Niveau de Risque'].str.contains('🚨', na=False, regex=False).sum()
                    warnings = df_config['Niveau de Risque'].str.contains('⚠️', na=False, regex=False).sum()
                    risques = critiques + warnings
                    ok = total - risques
                    
                    col1, col2, col3, col4 = st.columns(4)
                    col1.metric("📦 Total Programmes", total)
                    col2.metric("✅ OK", ok, delta=None if ok == total else f"-{risques}")
                    col3.metric("⚠️ Warnings", warnings)
                    col4.metric("🚨 CRITIQUES", critiques, delta_color="inverse")
            else:
                st.warning(T("Aucune donnée détectée", "No data detected"))
            
            st.markdown('</div>', unsafe_allow_html=True)

            # Résumé des risques
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("🔴 " + T("Résumé des Programmes à Risque", "Risk Programs Summary"))

            risk_summary = []
            if not df_config.empty and 'Niveau de Risque' in df_config.columns:
                for prog_name, group in df_config.groupby('Nom Programme'):
                    risks = group['Niveau de Risque'].unique()
                    has_risk = any('🚨' in str(r) or '⚠️' in str(r) for r in risks)
                    
                    if has_risk:
                        if any('🚨' in str(r) for r in risks):
                            env_couloir = group[['Environnement', 'Couloir']].drop_duplicates()
                            occurrences = len(group)
                            
                            if len(env_couloir) == 1:
                                env = env_couloir.iloc[0]['Environnement']
                                couloir = env_couloir.iloc[0]['Couloir']
                                risk_summary.append(
                                    f"🔴 {prog_name} : 🚨 Conflit de duplication dans le même environnement "
                                    f"({env} / {couloir}) - {occurrences} occurrence(s) détectée(s)"
                                )
                            else:
                                risk_summary.append(
                                    f"🔴 {prog_name} : 🚨 Conflits multiples détectés - {occurrences} occurrence(s)"
                                )
                        elif any('multi-environnements' in str(r) for r in risks):
                            envs = group['Environnement'].unique()
                            risk_summary.append(
                                f"🔴 {prog_name} : Existe dans {' et '.join(envs)} — "
                                f"Risque de modification inter-environnements"
                            )
                        elif any('multi-couloir' in str(r) for r in risks):
                            env = group['Environnement'].iloc[0]
                            couloirs = group['Couloir'].unique()
                            risk_summary.append(
                                f"🔴 {prog_name} : Programme présent dans {env} avec couloirs différents "
                                f"({', '.join(couloirs)})"
                            )

            if risk_summary:
                for risk in risk_summary:
                    if "🚨" in risk:
                        st.markdown(f'<div class="error-box" style="margin: 0.5rem 0;">{risk}</div>', unsafe_allow_html=True)
                    else:
                        st.markdown(f'<div class="warning-box" style="margin: 0.5rem 0;">{risk}</div>', unsafe_allow_html=True)
            else:
                st.success("✅ " + T("Aucun risque détecté", "No risk detected"))

            st.markdown('</div>', unsafe_allow_html=True)

            # Statistiques
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("📈 " + T("Statistiques", "Statistics"))

            if not df_config.empty and 'Niveau de Risque' in df_config.columns:
                total_programmes = len(df_config)
                programmes_uniques = df_config['Nom Programme'].nunique() if 'Nom Programme' in df_config.columns else total_programmes
                
                mask_critiques = df_config['Niveau de Risque'].str.contains('🚨', na=False, regex=False)
                conflits_critiques = mask_critiques.sum()
                programmes_critiques = df_config[mask_critiques]['Nom Programme'].nunique() if 'Nom Programme' in df_config.columns else conflits_critiques
                
                mask_warnings = df_config['Niveau de Risque'].str.contains('⚠️', na=False, regex=False) & ~mask_critiques
                warnings = mask_warnings.sum()
                programmes_warnings = df_config[mask_warnings]['Nom Programme'].nunique() if 'Nom Programme' in df_config.columns else warnings
                
                mask_ok = df_config['Niveau de Risque'].str.contains('—', na=False, regex=False) | df_config['Niveau de Risque'].isna()
                programmes_ok = df_config[mask_ok]['Nom Programme'].nunique() if 'Nom Programme' in df_config.columns else mask_ok.sum()
                
                programmes_a_risque = programmes_critiques + programmes_warnings
                
                # Affichage en colonnes
                col_main = st.columns(4)
                with col_main[0]:
                    st.metric("Total Lignes", total_programmes)
                with col_main[1]:
                    st.metric("Programmes Uniques", programmes_uniques)
                with col_main[2]:
                    st.metric("Programmes OK", programmes_ok, delta="✅" if programmes_ok > 0 else None)
                with col_main[3]:
                    st.metric("Programmes à Risque", programmes_a_risque, delta="⚠️" if programmes_a_risque > 0 else None, delta_color="inverse")
                
                st.markdown("---")
                
                col_risk = st.columns(3)
                with col_risk[0]:
                    st.metric("🚨 Conflits Critiques", programmes_critiques, delta="URGENT" if programmes_critiques > 0 else None, delta_color="inverse")
                with col_risk[1]:
                    st.metric("⚠️ Warnings", programmes_warnings, delta="À surveiller" if programmes_warnings > 0 else None, delta_color="inverse")
                with col_risk[2]:
                    taux_conformite = round((programmes_ok / programmes_uniques * 100) if programmes_uniques > 0 else 0, 1)
                    st.metric("📊 Taux de Conformité", f"{taux_conformite}%", delta="Bon" if taux_conformite >= 80 else "À améliorer", delta_color="normal" if taux_conformite >= 80 else "inverse")
            else:
                st.info(T("Aucune statistique disponible", "No statistics available"))

            st.markdown('</div>', unsafe_allow_html=True)

            # Bouton analyse avancée
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("🔍 " + T("Analyse Complémentaire", "Additional Analysis"))
            
            if st.button(
                "📊 " + T("Afficher/Masquer l'analyse approfondie", "Show/Hide in-depth analysis"),
                use_container_width=True,
                key="rgc_advanced_btn"
            ):
                st.session_state.show_advanced_analysis = not st.session_state.show_advanced_analysis

            st.markdown('</div>', unsafe_allow_html=True)

            # Analyse avancée (si demandée)
            if st.session_state.show_advanced_analysis and generate_metrics and not df_config.empty:
                st.markdown('<div class="glass-card">', unsafe_allow_html=True)
                st.subheader("📊 " + T("Analyses Avancées", "Advanced Analytics"))
                
                try:
                    import matplotlib.pyplot as plt
                    
                    fig, axes = plt.subplots(2, 2, figsize=(16, 10))
                    fig.patch.set_facecolor('#1E1E2E')
                    
                    # Graph 1
                    if 'Environnement' in df_config.columns:
                        env_counts = df_config['Environnement'].value_counts()
                        colors_env = ['#00D9A5', '#4EA3FF', '#FFB020']
                        axes[0, 0].pie(env_counts.values, labels=env_counts.index, autopct='%1.1f%%', colors=colors_env, textprops={'color': 'white', 'fontsize': 10})
                        axes[0, 0].set_title(T('Répartition par Environnement', 'Distribution by Environment'), color='white', fontsize=12, pad=20)
                        axes[0, 0].set_facecolor('#262636')
                    
                    # Graph 2
                    if 'Type de Programme' in df_config.columns:
                        lang_counts = df_config['Type de Programme'].value_counts()
                        axes[0, 1].bar(lang_counts.index, lang_counts.values, color=['#4EA3FF', '#FF6B9D', '#A3FFD6'])
                        axes[0, 1].set_title(T('Distribution par Langage', 'Distribution by Language'), color='white', fontsize=12, pad=20)
                        axes[0, 1].set_ylabel(T('Nombre de programmes', 'Number of programs'), color='white')
                        axes[0, 1].tick_params(colors='white')
                        axes[0, 1].set_facecolor('#262636')
                    
                    # Graph 3
                    if 'Modifié Par' in df_config.columns:
                        dev_counts = df_config['Modifié Par'].value_counts().head(10)
                        axes[1, 0].barh(dev_counts.index, dev_counts.values, color='#FF6B9D')
                        axes[1, 0].barh(dev_counts.index, dev_counts.values, color='#FF6B9D')
                        axes[1, 0].set_title(
                            T('Top 10 Développeurs Actifs', 'Top 10 Active Developers'),
                            color='white', fontsize=12, pad=20
                        )
                        axes[1, 0].set_xlabel(T('Nombre de modifications', 'Number of modifications'), color='white')
                        axes[1, 0].tick_params(colors='white')
                        axes[1, 0].set_facecolor('#262636')
                    
                    # Graph 4
                    if 'Niveau de Risque' in df_config.columns:
                        risk_categories = {
                            'OK': df_config['Niveau de Risque'].str.contains('—', na=False).sum(),
                            'Warning': df_config['Niveau de Risque'].str.contains('⚠️', na=False).sum() - df_config['Niveau de Risque'].str.contains('🚨', na=False).sum(),
                            'Critical': df_config['Niveau de Risque'].str.contains('🚨', na=False).sum()
                        }
                        
                        colors_risk = ['#00D9A5', '#FFB020', '#FF4757']
                        axes[1, 1].bar(
                            risk_categories.keys(),
                            risk_categories.values(),
                            color=colors_risk
                        )
                        axes[1, 1].set_title(
                            T('Répartition des Risques', 'Risk Distribution'),
                            color='white', fontsize=12, pad=20
                        )
                        axes[1, 1].set_ylabel(T('Nombre de programmes', 'Number of programs'), color='white')
                        axes[1, 1].tick_params(colors='white')
                        axes[1, 1].set_facecolor('#262636')
                    
                    plt.tight_layout()
                    st.pyplot(fig)
                    
                except ImportError:
                    st.info(T(
                        "📊 Installez matplotlib pour voir les graphiques : pip install matplotlib",
                        "📊 Install matplotlib to see charts: pip install matplotlib"
                    ))
                
                st.markdown('</div>', unsafe_allow_html=True)

            # Exports
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("💾 " + T("Exports", "Exports"))
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if not df_config.empty:
                    excel_buf = io.BytesIO()
                    try:
                        with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
                            df_config.to_excel(writer, index=False, sheet_name="Configuration_Complete")
                            
                            if 'rgc_df_summary' in st.session_state and st.session_state.rgc_df_summary is not None:
                                st.session_state.rgc_df_summary.to_excel(writer, index=False, sheet_name="Recap_Programmes")
                            
                            if 'Niveau de Risque' in df_config.columns:
                                df_risks = df_config[
                                    df_config['Niveau de Risque'].str.contains('⚠️|🚨', na=False)
                                ]
                                if not df_risks.empty:
                                    df_risks.to_excel(writer, index=False, sheet_name="Programmes_Risque")
                            
                            if st.session_state.rgc_stats_dict:
                                df_stats = pd.DataFrame(
                                    list(st.session_state.rgc_stats_dict.items()),
                                    columns=['Indicateur', 'Valeur']
                                )
                                df_stats.to_excel(writer, index=False, sheet_name="Statistiques")
                            
                            workbook = writer.book
                            header_format = workbook.add_format({
                                'bold': True,
                                'bg_color': '#4EA3FF',
                                'font_color': 'white',
                                'border': 1
                            })
                            
                            for sheet_name in writer.sheets:
                                worksheet = writer.sheets[sheet_name]
                                worksheet.set_row(0, 20, header_format)
                                worksheet.set_column(0, 20, 18)
                        
                        excel_buf.seek(0)
                        
                        st.download_button(
                            "📥 Excel Complet (4 feuilles)",
                            data=excel_buf,
                            file_name=f"RGC_Config_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True,
                            key="rgc_download_excel"
                        )
                    except ImportError:
                        st.info("💡 pip install xlsxwriter")
            
            with col2:
                if not df_config.empty:
                    csv_buf = io.StringIO()
                    df_config.to_csv(csv_buf, index=False, sep='|', encoding='utf-8')
                    csv_buf.seek(0)
                    
                    st.download_button(
                        "📄 CSV (Pipe)",
                        data=csv_buf.getvalue().encode('utf-8'),
                        file_name=f"RGC_Config_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        use_container_width=True,
                        key="rgc_download_csv"
                    )
            
            with col3:
                full_report = f"""
{'='*80}
RAPPORT D'ANALYSE RGC - GESTION DE CONFIGURATION
{'='*80}
Date        : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Fichier     : {st.session_state.rgc_uploaded_filename}
Analyste    : Système RGC IA

{'='*80}
TABLEAU DE CONFIGURATION
{'='*80}
{df_config.to_string() if not df_config.empty else 'Aucune donnée'}

{'='*80}
RÉSUMÉ DES RISQUES
{'='*80}
{chr(10).join(risk_summary) if risk_summary else 'Aucun risque détecté'}

{'='*80}
STATISTIQUES
{'='*80}
Total Programmes: {total_programmes if not df_config.empty else 0}
Programmes Uniques: {programmes_uniques if not df_config.empty else 0}
Programmes OK: {programmes_ok if not df_config.empty else 0}
Programmes à Risque: {programmes_a_risque if not df_config.empty else 0}
Conflits Critiques: {programmes_critiques if not df_config.empty else 0}
Warnings: {programmes_warnings if not df_config.empty else 0}

{'='*80}
FIN DU RAPPORT
{'='*80}
"""
                st.download_button(
                    "📋 Rapport TXT",
                    data=full_report.encode("utf-8"),
                    file_name=f"RGC_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    use_container_width=True,
                    key="rgc_download_txt"
                )
            
            st.markdown('</div>', unsafe_allow_html=True)

# ===================== MODE 5 : PSEUDO-CODE CONVERSION =====================
elif mode == TXT["modes"][0]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header("🔄 " + T(
        "Conversion Pseudo-Code → Langague Mainframe", 
        "Pseudo-Code → Mainframe Conversion language"
    ))
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="info-box">
        🎯 <strong>Mode Expert Mainframe </strong><br>
        Convertissez du pseudo-code en programmes mainframe :<br>
        • ✅ COBOL structuré et documenté<br>
        • ✅ PL/I optimisé<br>
        • ✅ Assembler (HLASM) <br>
        • ✅ Respect des standards IBM z/OS<br>
        • ✅ Code prêt à compilation
    </div>
    """, unsafe_allow_html=True)

    # Sélection du langage cible
    # Sélection du langage cible
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("⚙️ " + T("Configuration", "Settings"))
    
    col1, col2 = st.columns(2)
    with col1:
        target_lang = st.selectbox(
            "🎯 " + T("Langage cible", "Target Language"),
            ["COBOL", "PL/I", "HLASM (Assembler)"],
            help=T("Choisissez le langage mainframe de sortie", 
                   "Choose the output mainframe language"),
            key="pseudo_target_lang_select"  # ← CLÉ UNIQUE POUR LE WIDGET
        )
    
    with col2:
        program_name_input = st.text_input(
            "🏷️ " + T("Nom du programme", "Program Name"),
            value="PROGCONV",
            max_chars=8,
            help=T("Max 8 caractères alphanumériques", "Max 8 alphanumeric chars"),
            key="pseudo_program_name_input"  # ← CLÉ UNIQUE POUR LE WIDGET
        )
        
        # Sécuriser la conversion en majuscules
        program_name = program_name_input.upper() if program_name_input else "PROGCONV"
    
    st.markdown('</div>', unsafe_allow_html=True)

    # Zone de saisie du pseudo-code
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("📝 " + T("Pseudo-Code Source", "Source Pseudo-Code"))
    
    input_method = st.radio(
        T("Méthode de saisie", "Input Method"),
        [T("Saisie directe", "Direct Input"), T("Fichier texte", "Text File")],
        horizontal=True,
        key="pseudo_input_method_radio"  # ← CLÉ UNIQUE
    )
    
    pseudocode = ""
    
    if T("Saisie directe", "Direct Input") in input_method:
        # Utiliser la valeur stockée comme défaut
        default_value = st.session_state.pseudo_source_code if st.session_state.pseudo_source_code else ""
        
        pseudocode = st.text_area(
            T("Entrez votre pseudo-code", "Enter your pseudo-code"),
            placeholder=T(
                "Exemple :\n"
                "1. Lire fichier client FCLIENT\n"
                "2. Pour chaque enregistrement :\n"
                "   - Si solde > 1000 ALORS\n"
                "     * Calculer bonus = solde * 0.05\n"
                "     * Écrire dans FBONUS\n"
                "   - Sinon\n"
                "     * Incrémenter compteur clients standards\n"
                "3. Afficher total bonus calculés\n"
                "4. Fin programme",
                "Example:\n"
                "1. Read customer file FCLIENT\n"
                "2. For each record:\n"
                "   - If balance > 1000 THEN\n"
                "     * Calculate bonus = balance * 0.05\n"
                "     * Write to FBONUS\n"
                "   - Else\n"
                "     * Increment standard customer counter\n"
                "3. Display total bonuses calculated\n"
                "4. End program"
            ),
            height=300,
            key="pseudo_code_textarea_input",  # ← CLÉ UNIQUE
            value=default_value
        )
    else:
        uploaded_pseudo = st.file_uploader(
            "📂 " + T("Fichier pseudo-code (.txt)", "Pseudo-code file (.txt)"),
            type=["txt"],
            help=T("Fichier texte contenant le pseudo-code", "Text file with pseudo-code"),
            key="pseudo_file_uploader_widget"  # ← CLÉ UNIQUE
        )
        
        if uploaded_pseudo:
            try:
                pseudocode = uploaded_pseudo.read().decode("utf-8", errors="ignore")
                st.code(pseudocode[:500] + ("..." if len(pseudocode) > 500 else ""), 
                        language="text")
            except Exception as e:
                st.error(f"❌ Erreur lecture : {e}")
    
    st.markdown('</div>', unsafe_allow_html=True)

    # Validation du nom de programme
    valid_name = (
        len(program_name) > 0 and 
        len(program_name) <= 8 and 
        program_name.isalnum()
    )
    
    if not valid_name:
        st.markdown("""
        <div class="warning-box">
            ⚠️ Nom de programme invalide (1-8 caractères alphanumériques)
        </div>
        """, unsafe_allow_html=True)

    # Bouton de génération
    generate_button = st.button(
        "🚀 " + T("CONVERTIR EN " + target_lang, f"CONVERT TO {target_lang}"),
        disabled=not pseudocode.strip() or not valid_name,
        use_container_width=True,
        key="pseudo_convert_button"  # ← CLÉ UNIQUE
    )
    
    if generate_button:
        # STOCKER DANS SESSION_STATE avec des clés différentes des widgets
        st.session_state.pseudo_source_code = pseudocode
        st.session_state.pseudo_target_language = target_lang
        st.session_state.pseudo_generated_program_name = program_name
        
        # Déterminer la clé du prompt
        if "COBOL" in target_lang:
            lang_key = "COBOL"
        elif "PL/I" in target_lang or "PL1" in target_lang:
            lang_key = "PLI"
        elif "Assembler" in target_lang or "HLASM" in target_lang:
            lang_key = "HLASM"
        else:
            lang_key = "COBOL"
        
        # Charger le prompt depuis PromptEngine.yaml
        prompt_text = get_prompt(
            "PSEUDOCODE_CONVERSION",
            lang_key,
            pseudocode=pseudocode,
            program_name=program_name
        )
        
        # Vérifier si le prompt est vide ou contient des erreurs
        if not prompt_text or "{" in prompt_text:
            st.markdown(f"""
            <div class="error-box">
                ❌ Erreur de configuration du prompt pour {target_lang}<br>
                Vérifiez que <code>PromptEngine.yaml</code> contient la section 
                <code>PSEUDOCODE_CONVERSION.{lang_key}</code>
            </div>
            """, unsafe_allow_html=True)
            st.stop()
        
        client = llm_client(max_tokens=4000, temperature=0.1)
        
        if not client:
            st.markdown('<div class="error-box">❌ Client LLM indisponible</div>', 
                        unsafe_allow_html=True)
        else:
            with st.spinner(T(
                f"🧠 Génération du code {target_lang}...",
                f"🧠 Generating {target_lang} code..."
            )):
                try:
                    response = client.invoke(prompt_text)
                    result = response.content if hasattr(response, 'content') else str(response)
                    
                    if result:
                        # Nettoyage du résultat
                        code_clean = result.strip()
                        
                        # Supprimer les balises markdown
                        if code_clean.startswith("```"):
                            lines = code_clean.split("\n")
                            if lines[-1].strip() == "```":
                                lines = lines[1:-1]
                            else:
                                lines = lines[1:]
                            code_clean = "\n".join(lines)
                        
                        # Stocker dans session_state
                        st.session_state.pseudo_result = code_clean
                        
                except Exception as e:
                    st.error(f"❌ Erreur lors de la génération : {e}")
                    import traceback
                    st.code(traceback.format_exc(), language="python")
                    st.session_state.pseudo_result = None

    # Affichage des résultats (persistant)
    if st.session_state.pseudo_result:
        code_clean = st.session_state.pseudo_result
        target_lang = st.session_state.pseudo_target_language
        program_name = st.session_state.pseudo_generated_program_name
        
        # Affichage du code généré
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader(f"📘 Code {target_lang} Généré")
        
        # Choisir le langage de coloration syntaxique
        if "COBOL" in target_lang:
            syntax_lang = "cobol"
        elif "PL/I" in target_lang:
            syntax_lang = "sql"
        elif "Assembler" in target_lang:
            syntax_lang = "asm"
        else:
            syntax_lang = "text"
        
        st.code(code_clean, language=syntax_lang)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Exports
        col1, col2 = st.columns(2)
        
        # Déterminer l'extension
        ext_map = {
            "COBOL": ".cbl",
            "PL/I": ".pli",
            "HLASM (Assembler)": ".asm"
        }
        file_ext = ext_map.get(target_lang, ".txt")
        
        with col1:
            st.download_button(
                f"💾 Télécharger {program_name}{file_ext}",
                data=code_clean.encode("utf-8"),
                file_name=f"{program_name}{file_ext}",
                use_container_width=True,
                key="pseudo_download_code_btn"  # ← CLÉ UNIQUE
            )
        
        with col2:
            # Export avec documentation complète
            doc_complete = f"""{'='*80}
PROGRAMME MAINFRAME GÉNÉRÉ PAR IA
{'='*80}
Langage      : {target_lang}
Programme    : {program_name}
Date         : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Généré par   : Assistant IA Mainframe Expert

{'='*80}
PSEUDO-CODE SOURCE
{'='*80}
{st.session_state.pseudo_source_code}

{'='*80}
CODE {target_lang} GÉNÉRÉ
{'='*80}
{code_clean}

{'='*80}
FIN DU DOCUMENT
{'='*80}
"""
            st.download_button(
                "📄 Documentation complète (.txt)",
                data=doc_complete.encode("utf-8"),
                file_name=f"{program_name}_DOCUMENTATION.txt",
                use_container_width=True,
                key="pseudo_download_doc_btn"  # ← CLÉ UNIQUE
            )
        
        # Métriques du code généré
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("📊 " + T("Métriques du code", "Code Metrics"))
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Lignes de code", len(code_clean.split('\n')))
        with col2:
            st.metric("Caractères", len(code_clean))
        with col3:
            # Compter les commentaires selon le langage
            if "COBOL" in target_lang:
                comment_lines = sum(1 for line in code_clean.split('\n') 
                                   if line.strip().startswith('*') or '*>' in line)
            elif "PL/I" in target_lang:
                comment_lines = code_clean.count('/*')
            else:  # Assembler
                comment_lines = sum(1 for line in code_clean.split('\n') 
                                   if line.strip().startswith('*'))
            
            st.metric("Lignes commentées", comment_lines)
        
        st.markdown('</div>', unsafe_allow_html=True)
# ===================== MODE 6 : EXTRACTION RÈGLES DE GESTION =====================
elif mode == TXT["modes"][5]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header("📊 " + T(
        "Extraction de Règles de Gestion - Code Legacy", 
        "Business Rules Extraction - Legacy Code"
    ))
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="info-box">
        🎯 <strong>Mode Expert Analyse Legacy</strong><br>
        Extrayez les règles de gestion depuis du code mainframe :<br>
        • ✅ Analyse complète COBOL/PL/I/Assembler<br>
        • ✅ Dictionnaire de données automatique<br>
        • ✅ Règles numérotées et documentées<br>
        • ✅ Scénarios de tests recommandés<br>
        • ✅ Document lisible par non-techniciens<br>
        <br>
        <strong>Format de sortie :</strong> Document structuré (Word + TXT)
    </div>
    """, unsafe_allow_html=True)

    # Upload du fichier source
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("📂 " + T("Fichier Source Legacy", "Legacy Source File"))
    
    uploaded_legacy = st.file_uploader(
        "📄 " + T("Programme COBOL/PL/I/ASM", "COBOL/PL/I/ASM Program"),
        type=["cbl", "cob", "pli", "pl1", "asm", "txt"],
        help=T("Sélectionnez votre fichier source mainframe", "Select your mainframe source file"),
        key="business_rules_file_uploader"
    )
    
    if uploaded_legacy:
        st.markdown(f"""
        <div class="success-box">
            ✅ Fichier chargé : <strong>{uploaded_legacy.name}</strong>
        </div>
        """, unsafe_allow_html=True)
        
        st.session_state.business_rules_filename = uploaded_legacy.name
        
        # Détection du langage
        file_name = uploaded_legacy.name
        ext = file_name.split(".")[-1].lower()
        
        lang_map = {
            "cbl": "COBOL", 
            "cob": "COBOL", 
            "pli": "PL/I", 
            "pl1": "PL/I", 
            "asm": "ASM",
            "txt": "COBOL"  # Par défaut
        }
        detected_lang = lang_map.get(ext, "COBOL")
        
        st.markdown(f"""
        <div class="info-box">
            🔍 Langage détecté : <strong>{detected_lang}</strong>
        </div>
        """, unsafe_allow_html=True)
        
        st.session_state.business_rules_language = detected_lang
        
        # Lecture du code
        try:
            uploaded_legacy.seek(0)
            source_code = uploaded_legacy.read().decode("utf-8", errors="ignore")
            st.session_state.business_rules_source_code = source_code
            
            # Aperçu du code
            st.markdown("**📋 Aperçu du code source :**")
            preview_lines = source_code.split('\n')[:20]
            st.code('\n'.join(preview_lines), language="cobol" if "COBOL" in detected_lang else "text")
            
            if len(source_code.split('\n')) > 20:
                st.caption(f"... ({len(source_code.split('\n')) - 20} lignes supplémentaires)")
        
        except Exception as e:
            st.error(f"❌ Erreur de lecture : {e}")
            st.stop()
    
    st.markdown('</div>', unsafe_allow_html=True)

    # Options d'extraction
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("⚙️ " + T("Options d'extraction", "Extraction Options"))
    
    col1, col2, col3 = st.columns(3)
    with col1:
        extract_dict = st.checkbox(
            "📖 Dictionnaire de données",
            value=True,
            help=T("Générer le dictionnaire complet des variables", "Generate complete data dictionary"),
            key="br_extract_dict"
        )
    with col2:
        extract_tests = st.checkbox(
            "🧪 Scénarios de tests",
            value=True,
            help=T("Proposer des scénarios de test", "Suggest test scenarios"),
            key="br_extract_tests"
        )
    with col3:
        extract_diagrams = st.checkbox(
            "📊 Suggestion diagrammes",
            value=False,
            help=T("Suggérer des diagrammes de flux", "Suggest flow diagrams"),
            key="br_extract_diagrams"
        )
    
    st.markdown('</div>', unsafe_allow_html=True)

    # Bouton d'extraction
    extract_button = st.button(
        "🚀 " + T("EXTRAIRE LES RÈGLES DE GESTION", "EXTRACT BUSINESS RULES"),
        disabled=not uploaded_legacy,
        use_container_width=True,
        key="business_rules_extract_btn"
    )

    if extract_button:
        source_code = st.session_state.business_rules_source_code
        detected_lang = st.session_state.business_rules_language
        
        if not source_code or not source_code.strip():
            st.markdown('<div class="error-box">❌ Code source vide</div>', unsafe_allow_html=True)
            st.stop()
        
        # Déterminer la clé du prompt
        lang_key = "COBOL"
        if "PL/I" in detected_lang or "PL1" in detected_lang:
            lang_key = "PLI"
        elif "ASM" in detected_lang or "Assembler" in detected_lang:
            lang_key = "ASM"
        
        # Charger le prompt
        prompt_text = get_prompt(
            "BUSINESS_RULES_EXTRACTION",
            lang_key,
            source_code=source_code[:100_000]  # Limiter à 100k caractères
        )
        
        if not prompt_text or "{" in prompt_text:
            st.markdown(f"""
            <div class="error-box">
                ❌ Erreur de configuration du prompt pour {detected_lang}<br>
                Vérifiez que <code>PromptEngine.yaml</code> contient la section 
                <code>BUSINESS_RULES_EXTRACTION.{lang_key}</code>
            </div>
            """, unsafe_allow_html=True)
            st.stop()
        
        client = llm_client(max_tokens=4000, temperature=0.1)
        
        if not client:
            st.markdown('<div class="error-box">❌ Client LLM indisponible</div>', unsafe_allow_html=True)
            st.stop()
        
        with st.spinner(T(
            "🧠 Extraction des règles de gestion en cours...",
            "🧠 Extracting business rules..."
        )):
            try:
                response = client.invoke(prompt_text)
                result = response.content if hasattr(response, 'content') else str(response)
                st.session_state.business_rules_result = result
            except Exception as e:
                st.error(f"❌ Erreur LLM : {e}")
                import traceback
                st.code(traceback.format_exc(), language="python")
                st.session_state.business_rules_result = None

    # Affichage des résultats (persistant)
    if st.session_state.business_rules_result:
        result = st.session_state.business_rules_result
        detected_lang = st.session_state.business_rules_language
        filename = st.session_state.business_rules_filename
        
        # Parsing des sections
        sections = {
            'resume': '',
            'entrees_sorties': '',
            'dictionnaire': '',
            'regles': '',
            'scenarios': '',
            'ambiguites': ''
        }
        
        if "=== RESUME_FONCTIONNEL ===" in result:
            parts = result.split("=== RESUME_FONCTIONNEL ===", 1)[1]
            
            if "=== ENTREES_SORTIES ===" in parts:
                sections['resume'], rest = parts.split("=== ENTREES_SORTIES ===", 1)
                
                if "=== DICTIONNAIRE_DONNEES ===" in rest:
                    sections['entrees_sorties'], rest = rest.split("=== DICTIONNAIRE_DONNEES ===", 1)
                    
                    if "=== REGLES_GESTION ===" in rest:
                        sections['dictionnaire'], rest = rest.split("=== REGLES_GESTION ===", 1)
                        
                        if "=== SCENARIOS_TESTS ===" in rest:
                            sections['regles'], rest = rest.split("=== SCENARIOS_TESTS ===", 1)
                            
                            if "=== POINTS_AMBIGUITE ===" in rest:
                                sections['scenarios'], sections['ambiguites'] = rest.split("=== POINTS_AMBIGUITE ===", 1)
                            else:
                                sections['scenarios'] = rest
                        else:
                            sections['regles'] = rest

        # Affichage organisé par onglets
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "📋 Résumé Fonctionnel",
            "📖 Dictionnaire Données",
            "⚖️ Règles de Gestion",
            "🧪 Scénarios Tests",
            "⚠️ Points d'Ambiguïté"
        ])
        
        with tab1:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("📋 Résumé Fonctionnel")
            st.markdown(sections['resume'] if sections['resume'] else result[:1000])
            st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("📥 Entrées / Sorties")
            st.markdown(sections['entrees_sorties'] if sections['entrees_sorties'] else "_Non détecté_")
            st.markdown('</div>', unsafe_allow_html=True)
        with tab2:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("📖 Dictionnaire des Données")
            
            if sections['dictionnaire']:
                # Essayer de parser en DataFrame
                dict_lines = sections['dictionnaire'].strip().split('\n')
                dict_data = []
                headers = []
                
                for line in dict_lines:
                    if '|' in line and not line.strip().startswith('-'):
                        parts = [p.strip() for p in line.split('|')]
                        
                        # Filtrer les éléments vides au début/fin
                        parts = [p for p in parts if p]
                        
                        if len(parts) >= 4:
                            if not headers:
                                # Première ligne avec des headers valides
                                headers = parts
                                # Vérifier qu'il n'y a pas de doublons
                                if len(headers) != len(set(headers)):
                                    # Renommer les doublons
                                    seen = {}
                                    new_headers = []
                                    for h in headers:
                                        if h in seen:
                                            seen[h] += 1
                                            new_headers.append(f"{h}_{seen[h]}")
                                        else:
                                            seen[h] = 0
                                            new_headers.append(h)
                                    headers = new_headers
                            elif len(parts) == len(headers):
                                dict_data.append(parts)
                
                if headers and dict_data:
                    try:
                        df_dict = pd.DataFrame(dict_data, columns=headers)
                        st.dataframe(df_dict, use_container_width=True)
                    except Exception as e:
                        st.warning(f"⚠️ Impossible de créer le tableau : {e}")
                        st.markdown(sections['dictionnaire'])
                else:
                    st.markdown(sections['dictionnaire'])
            else:
                st.info("Aucun dictionnaire détecté")
            
            st.markdown('</div>', unsafe_allow_html=True)
            
        with tab3:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("⚖️ Règles de Gestion Extraites")
            
            if sections['regles']:
                # Compter les règles
                nb_regles = sections['regles'].count('**Règle RG-')
                st.metric("📊 Nombre de règles détectées", nb_regles)
                st.markdown("---")
                st.markdown(sections['regles'])
            else:
                st.warning("Aucune règle extraite")
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        with tab4:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("🧪 Scénarios de Tests Recommandés")
            st.markdown(sections['scenarios'] if sections['scenarios'] else "_Aucun scénario proposé_")
            st.markdown('</div>', unsafe_allow_html=True)
        
        with tab5:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("⚠️ Points d'Ambiguïté ou Informations Manquantes")
            
            if sections['ambiguites'].strip():
                st.warning(sections['ambiguites'])
            else:
                st.success("✅ Aucune ambiguïté détectée")
            
            st.markdown('</div>', unsafe_allow_html=True)

        # Exports
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("💾 " + T("Exports Documentation", "Documentation Exports"))
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Export Word (si disponible)
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Titre
                title = doc.add_heading(f"Documentation Règles de Gestion - {filename}", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Métadonnées
                doc.add_paragraph(f"Langage : {detected_lang}")
                doc.add_paragraph(f"Date d'extraction : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Généré par : Assistant IA Mainframe Expert")
                doc.add_paragraph("="*80)
                
                # Sections
                if sections['resume']:
                    doc.add_heading("1. Résumé Fonctionnel", 1)
                    doc.add_paragraph(sections['resume'])
                
                if sections['entrees_sorties']:
                    doc.add_heading("2. Entrées / Sorties", 1)
                    doc.add_paragraph(sections['entrees_sorties'])
                
                if sections['dictionnaire']:
                    doc.add_heading("3. Dictionnaire des Données", 1)
                    doc.add_paragraph(sections['dictionnaire'])
                
                if sections['regles']:
                    doc.add_heading("4. Règles de Gestion", 1)
                    doc.add_paragraph(sections['regles'])
                
                if sections['scenarios']:
                    doc.add_heading("5. Scénarios de Tests", 1)
                    doc.add_paragraph(sections['scenarios'])
                
                if sections['ambiguites']:
                    doc.add_heading("6. Points d'Ambiguïté", 1)
                    doc.add_paragraph(sections['ambiguites'])
                
                # Sauvegarde
                doc_buf = io.BytesIO()
                doc.save(doc_buf)
                doc_buf.seek(0)
                
                st.download_button(
                    "📥 Télécharger Documentation (Word)",
                    data=doc_buf,
                    file_name=f"Regles_Gestion_{filename.split('.')[0]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="br_download_word"
                )
            except ImportError:
                st.info("💡 Installez python-docx pour l'export Word")
        
        with col2:
            # Export TXT complet
            txt_export = f"""
{'='*80}
DOCUMENTATION RÈGLES DE GESTION
{'='*80}
Fichier      : {filename}
Langage      : {detected_lang}
Date         : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Généré par   : Assistant IA Mainframe Expert

{'='*80}
1. RÉSUMÉ FONCTIONNEL
{'='*80}
{sections['resume']}

{'='*80}
2. ENTRÉES / SORTIES
{'='*80}
{sections['entrees_sorties']}

{'='*80}
3. DICTIONNAIRE DES DONNÉES
{'='*80}
{sections['dictionnaire']}

{'='*80}
4. RÈGLES DE GESTION
{'='*80}
{sections['regles']}

{'='*80}
5. SCÉNARIOS DE TESTS
{'='*80}
{sections['scenarios']}

{'='*80}
6. POINTS D'AMBIGUÏTÉ
{'='*80}
{sections['ambiguites']}

{'='*80}
FIN DU DOCUMENT
{'='*80}
"""
            st.download_button(
                "📄 Télécharger Documentation (TXT)",
                data=txt_export.encode("utf-8"),
                file_name=f"Regles_Gestion_{filename.split('.')[0]}.txt",
                use_container_width=True,
                key="br_download_txt"
            )
        
        st.markdown('</div>', unsafe_allow_html=True)

        # Métriques finales
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("📊 " + T("Statistiques d'extraction", "Extraction Statistics"))
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            nb_regles = sections['regles'].count('**Règle RG-')
            st.metric("Règles extraites", nb_regles)
        with col2:
            nb_variables = sections['dictionnaire'].count('|') // 4 if sections['dictionnaire'] else 0
            st.metric("Variables documentées", nb_variables)
        with col3:
            nb_scenarios = sections['scenarios'].count('Cas normal')
            st.metric("Scénarios tests", nb_scenarios)
        with col4:
            nb_ambiguites = sections['ambiguites'].count('-')
            st.metric("Points d'ambiguïté", nb_ambiguites)
        
        st.markdown('</div>', unsafe_allow_html=True)
# ===================== MODE 7 : APPLICATION ANALYZER =====================
# ===================== MODE 7 : APPLICATION ANALYZER PRO =====================
# ===================== MODE 7 : APPLICATION ANALYZER ULTIMATE =====================
elif mode == TXT["modes"][6]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header("🔍 " + T(
        "Application Analyzer ULTIMATE - Reverse Engineering", 
        "Application Analyzer ULTIMATE - Reverse Engineering"
    ))
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="info-box">
        🎯 <strong>Analyse Complète Automatisée par IA</strong><br>
        <br>
        <strong>🔬 Parsing Statique Avancé :</strong><br>
        • Détection automatique : COBOL Batch/CICS/IMS/DB2/MQ/Hybrid<br>
        • Extraction appels : CALL, CICS LINK/XCTL, IMS DL/I, EXEC SQL, MQ<br>
        • Identification ressources : Transactions, Tables, Segments, Queues<br>
        • Analyse copybooks, fichiers VSAM, JCL steps<br>
        <br>
        <strong>🤖 Analyse Intelligente :</strong><br>
        • Compréhension du rôle métier de chaque programme<br>
        • Explication du flux de données en langage naturel<br>
        • Identification des points critiques et risques<br>
        • Recommandations de modernisation personnalisées<br>
        • Génération de documentation technique + métier<br>
        <br>
        <strong>📊 Livrables :</strong><br>
        • Graphe de dépendances interactif avec légende<br>
        • Rapport d'analyse détaillé par programme (IA)<br>
        • Matrice d'impact (qui appelle qui, qui modifie quoi)<br>
        • Plan de migration priorisé<br>
        • Export JSON, Excel multi-feuilles, Markdown, ZIP<br>
        <br>
        <strong>Format attendu :</strong> ZIP (.cbl, .cob, .jcl, .cpy, .bms, .txt)
    </div>
    """, unsafe_allow_html=True)

    # ========== UPLOAD ZIP ==========
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("📂 " + T("Chargement du ZIP source", "Source ZIP Upload"))
    
    uploaded_zip_ultimate = st.file_uploader(
        "📦 " + T("Fichier ZIP (Application Mainframe)", "ZIP file (Mainframe Application)"),
        type=["zip"],
        help=T(
            "ZIP contenant programmes COBOL, JCL, copybooks, BMS...",
            "ZIP containing COBOL programs, JCL, copybooks, BMS..."
        ),
        key="ultimate_analyzer_zip_uploader"
    )
    
    if uploaded_zip_ultimate:
        file_size_mb = uploaded_zip_ultimate.size / (1024 * 1024)
        st.markdown(f"""
        <div class="success-box">
            ✅ <strong>ZIP chargé : {uploaded_zip_ultimate.name}</strong> ({file_size_mb:.2f} MB)
        </div>
        """, unsafe_allow_html=True)
        
        # Extraction
        try:
            with zipfile.ZipFile(uploaded_zip_ultimate, "r") as zip_ref:
                file_list = zip_ref.namelist()
                
                valid_files = []
                skipped_files = []
                
                for filename in file_list:
                    if filename.endswith('/') or filename.startswith(('__MACOSX', '.', '_')):
                        skipped_files.append(filename)
                        continue
                    
                    if any(filename.lower().endswith(ext) for ext in ['.cbl', '.cob', '.jcl', '.txt', '.cpy', '.copy', '.bms', '.asm']):
                        try:
                            content = zip_ref.read(filename)
                            valid_files.append((filename, content))
                        except Exception as e:
                            st.warning(f"⚠️ {filename}: {e}")
                    else:
                        skipped_files.append(filename)
                
                st.session_state.analyzer_uploaded_files = valid_files
                
                if valid_files:
                    st.markdown(f"""
                    <div class="success-box">
                        ✅ <strong>{len(valid_files)} fichier(s) prêt(s) à analyser</strong>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Aperçu
                    file_types = defaultdict(list)
                    for fname, _ in valid_files:
                        basename = fname.split('/')[-1].lower()
                        if basename.endswith(('.cbl', '.cob')):
                            file_types['COBOL'].append(fname)
                        elif basename.endswith('.jcl'):
                            file_types['JCL'].append(fname)
                        elif basename.endswith(('.cpy', '.copy')):
                            file_types['COPYBOOK'].append(fname)
                        elif basename.endswith('.bms'):
                            file_types['BMS'].append(fname)
                        else:
                            file_types['AUTRE'].append(fname)
                    
                    col1, col2, col3, col4, col5 = st.columns(5)
                    col1.metric("📘 COBOL", len(file_types.get('COBOL', [])))
                    col2.metric("🔧 JCL", len(file_types.get('JCL', [])))
                    col3.metric("📖 COPYBOOK", len(file_types.get('COPYBOOK', [])))
                    col4.metric("🖥️ BMS", len(file_types.get('BMS', [])))
                    col5.metric("📄 AUTRE", len(file_types.get('AUTRE', [])))
                    
                    with st.expander("🔍 Détail des fichiers"):
                        for ftype, files in sorted(file_types.items()):
                            if files:
                                st.markdown(f"**{ftype} ({len(files)}) :**")
                                for f in sorted(files)[:10]:
                                    st.text(f"  • {f}")
                                if len(files) > 10:
                                    st.caption(f"... et {len(files) - 10} autres")
                                st.markdown("---")
                else:
                    st.warning("⚠️ Aucun fichier valide trouvé")
        
        except zipfile.BadZipFile:
            st.error("❌ ZIP invalide")
        except Exception as e:
            st.error(f"❌ Erreur : {e}")
    
    st.markdown('</div>', unsafe_allow_html=True)

    # ========== OPTIONS D'ANALYSE ==========
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("⚙️ " + T("Options d'analyse", "Analysis Options"))
    
    col_opt1, col_opt2, col_opt3 = st.columns(3)
    
    with col_opt1:
        enable_llm_analysis = st.checkbox(
            "🤖 Analyse IA (LLM)",
            value=True,
            help=T(
                "Analyse intelligente du rôle métier de chaque programme",
                "Intelligent business role analysis for each program"
            ),
            key="ultimate_enable_llm_checkbox"
        )
        
        detect_orphans_ult = st.checkbox(
            "🚫 Détecter orphelins",
            value=True,
            key="ultimate_orphans_checkbox"
        )
    
    with col_opt2:
        analyze_dependencies = st.checkbox(
            "🔗 Graphe dépendances",
            value=True,
            help=T("Graphe complet avec ressources", "Complete graph with resources"),
            key="ultimate_deps_checkbox"
        )
        
        detect_critical_ult = st.checkbox(
            "⚠️ Programmes critiques",
            value=True,
            key="ultimate_critical_checkbox"
        )
    
    with col_opt3:
        generate_migration_plan = st.checkbox(
            "🚀 Plan migration",
            value=True,
            help=T("Recommandations de modernisation", "Modernization recommendations"),
            key="ultimate_migration_checkbox"
        )
        
        detect_cycles_ult = st.checkbox(
            "🔄 Cycles dépendances",
            value=True,
            key="ultimate_cycles_checkbox"
        )
    
    st.markdown('</div>', unsafe_allow_html=True)

    # ========== BOUTON ANALYSE ==========
    analyze_ultimate_button = st.button(
        "🚀 " + T("LANCER ANALYSE COMPLÈTE", "START COMPLETE ANALYSIS"),
        disabled=not uploaded_zip_ultimate or not st.session_state.analyzer_uploaded_files,
        use_container_width=True,
        key="ultimate_analyze_btn"
    )

    # ========== ANALYSE PRINCIPALE ==========
    if analyze_ultimate_button and st.session_state.analyzer_uploaded_files:
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Étape 1 : Parsing statique
            status_text.text("🔧 Phase 1/5 : Parsing statique des sources...")
            progress_bar.progress(10)
            
            analyzer = MainframeReverseEngineer()
            analyzer.analyze_sources(st.session_state.analyzer_uploaded_files)
            
            progress_bar.progress(25)
            
            # Étape 2 : Statistiques
            status_text.text("📊 Phase 2/5 : Calcul des statistiques...")
            stats = analyzer.compute_statistics()
            
            progress_bar.progress(40)
            
            # Étape 3 : Graphe
            status_text.text("🔗 Phase 3/5 : Construction du graphe de dépendances...")
            graph = analyzer.build_dependency_graph() if analyze_dependencies else None
            
            progress_bar.progress(55)
            
            # Étape 4 : Cycles
            cycles = []
            if detect_cycles_ult and graph:
                try:
                    import networkx as nx
                    cycles = list(nx.simple_cycles(graph))
                except:
                    pass
            
            progress_bar.progress(65)
            
            # Étape 5 : Analyse LLM (NOUVELLE PARTIE)
            llm_analyses = {}
            
            if enable_llm_analysis:
                status_text.text("🤖 Phase 4/5 : Analyse intelligente par IA...")
                
                client = llm_client(max_tokens=1500, temperature=0.2)
                
                if client:
                    # Limiter à 10 programmes pour la démo (sinon trop long)
                    programs_to_analyze = list(analyzer.programs.values())[:10]
                    
                    for idx, prog in enumerate(programs_to_analyze, 1):
                        status_text.text(f"🤖 Analyse IA : {prog.name} ({idx}/{len(programs_to_analyze)})...")
                        
                        # Préparer le contexte pour le LLM
                        prog_context = f"""
**Programme : {prog.name}**
**Type : {prog.component_type.value}**
**Lignes : {prog.lines}**
**Complexité : {prog.complexity_score}**

**Appels émis :** {', '.join(prog.calls_static[:5]) if prog.calls_static else 'Aucun'}
**Appelé par :** {', '.join(prog.called_by[:5]) if prog.called_by else 'Aucun'}

**CICS :** {'Oui - ' + str(len(prog.calls_cics)) + ' appels' if prog.calls_cics else 'Non'}
**DB2 Tables :** {', '.join(prog.db2_tables[:3]) if prog.db2_tables else 'Aucune'}
**IMS Segments :** {', '.join(prog.ims_segments[:3]) if prog.ims_segments else 'Aucun'}
**MQ Queues :** {', '.join(prog.mq_queues) if prog.mq_queues else 'Aucune'}

**Transactions CICS :** {', '.join(prog.transactions) if prog.transactions else 'Aucune'}
**BMS Maps :** {', '.join(prog.bms_maps) if prog.bms_maps else 'Aucun'}
**Copybooks :** {', '.join(prog.copybooks[:5]) if prog.copybooks else 'Aucun'}
"""
                        
                        prompt_llm = f"""Tu es un expert en analyse d'applications Mainframe.

Analyse ce programme et fournis une réponse structurée en français :

{prog_context}

**Réponds uniquement avec ce format (sans ajouter de texte avant ou après) :**

RÔLE MÉTIER: [Une phrase décrivant le rôle métier du programme]

FLUX DE DONNÉES: [Description du flux : quelles données en entrée, quelles données en sortie, transformations]

DÉPENDANCES CRITIQUES: [Liste des dépendances importantes et pourquoi elles sont critiques]

NIVEAU DE RISQUE: [FAIBLE/MOYEN/ÉLEVÉ avec justification courte]

RECOMMANDATIONS: [2-3 recommandations concrètes pour amélioration ou migration]
"""
                        
                        try:
                            response = client.invoke(prompt_llm)
                            llm_result = response.content if hasattr(response, 'content') else str(response)
                            llm_analyses[prog.name] = llm_result
                        except Exception as e:
                            llm_analyses[prog.name] = f"Erreur analyse IA : {e}"
                        
                        # Mettre à jour la progress bar
                        progress_increment = (idx / len(programs_to_analyze)) * 20
                        progress_bar.progress(int(65 + progress_increment))
                else:
                    st.warning("⚠️ Client LLM indisponible, analyse statique uniquement")
            
            progress_bar.progress(85)
            
            # Étape 6 : Plan de migration (avec LLM si activé)
            migration_plan = None
            
            if generate_migration_plan:
                status_text.text("🚀 Phase 5/5 : Génération du plan de migration...")
                
                if enable_llm_analysis and client:
                    migration_prompt = f"""Tu es un architecte expert en modernisation Mainframe.

Voici le contexte de l'application analysée :

- **{stats['total_programs']} programmes** au total
- **{stats['cobol_cics']} programmes CICS** (transactionnels)
- **{stats['cobol_batch']} programmes Batch**
- **{stats['total_db2_tables']} tables DB2**
- **{stats['total_transactions']} transactions CICS**
- **{stats['orphan_programs']} programmes orphelins** (non utilisés)
- **{stats['critical_programs']} programmes critiques** (forte dépendance)

Propose un **plan de migration en 3 phases** (court/moyen/long terme) avec actions concrètes, bénéfices attendus et risques.

Format attendu :

PHASE 1 - COURT TERME (0-6 mois):
- Action 1: [description]
- Action 2: [description]
Bénéfices: [liste]
Risques: [liste]

PHASE 2 - MOYEN TERME (6-18 mois):
...

PHASE 3 - LONG TERME (18-36 mois):
...

RECOMMANDATIONS PRIORITAIRES:
1. [recommandation]
2. [recommandation]
"""
                    
                    try:
                        response = client.invoke(migration_prompt)
                        migration_plan = response.content if hasattr(response, 'content') else str(response)
                    except:
                        migration_plan = "Plan de migration non disponible"
            
            progress_bar.progress(95)
            
            # Rapport JSON
            status_text.text("📄 Génération du rapport final...")
            report_json = analyzer.generate_report_json()
            
            # Ajouter les analyses LLM au rapport
            if llm_analyses:
                report_json['llm_analyses'] = llm_analyses
            
            if migration_plan:
                report_json['migration_plan'] = migration_plan
            
            # Stockage
            st.session_state.analyzer_results = analyzer
            st.session_state.analyzer_computed_metrics = stats
            st.session_state.analyzer_dependency_graph = graph
            st.session_state.analyzer_generated_report = report_json
            st.session_state.analyzer_detected_cycles = cycles
            st.session_state.analyzer_llm_analyses = llm_analyses
            st.session_state.analyzer_migration_plan = migration_plan
            
            progress_bar.progress(100)
            status_text.empty()
            progress_bar.empty()
            
            st.success("✅ " + T("Analyse terminée avec succès !", "Analysis completed successfully!"))
            
            # Résumé
            col_sum1, col_sum2, col_sum3, col_sum4 = st.columns(4)
            col_sum1.metric("📦 Programmes", stats['total_programs'])
            col_sum2.metric("🤖 Analyses IA", len(llm_analyses))
            col_sum3.metric("⚠️ Critiques", stats['critical_programs'])
            col_sum4.metric("🔄 Cycles", len(cycles))
        
        except Exception as e:
            st.error(f"❌ Erreur : {e}")
            import traceback
            st.code(traceback.format_exc(), language="python")

    # ========== AFFICHAGE RÉSULTATS ==========
    if st.session_state.analyzer_results:
        analyzer = st.session_state.analyzer_results
        stats = st.session_state.analyzer_computed_metrics
        graph = st.session_state.analyzer_dependency_graph
        report_json = st.session_state.analyzer_generated_report
        cycles = st.session_state.get('analyzer_detected_cycles', [])
        llm_analyses = st.session_state.get('analyzer_llm_analyses', {})
        migration_plan = st.session_state.get('analyzer_migration_plan', None)

        # ONGLETS
        tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
            "📊 Dashboard",
            "📘 Programmes",
            "🤖 Analyses IA",
            "🔗 Graphe",
            "🔍 Impact",
            "⚠️ Alertes",
            "🚀 Migration",
            "💾 Exports"
        ])

        # TAB 1 : DASHBOARD
        with tab1:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("📊 Vue d'Ensemble")
            
            col1, col2, col3, col4, col5 = st.columns(5)
            col1.metric("📦 Programmes", stats['total_programs'])
            col2.metric("📝 Lignes", f"{stats['total_lines']:,}")
            col3.metric("🔧 Jobs JCL", stats['total_jcl'])
            col4.metric("📖 Copybooks", stats['total_copybooks'])
            col5.metric("🧮 Complexité", f"{stats['avg_complexity']:.1f}")
            
            st.markdown("---")
            st.markdown("**🏗️ Répartition par type**")
            
            col6, col7, col8, col9, col10 = st.columns(5)
            col6.metric("Batch", stats['cobol_batch'])
            col7.metric("CICS", stats['cobol_cics'])
            col8.metric("IMS", stats['cobol_ims'])
            col9.metric("DB2", stats['cobol_db2'])
            col10.metric("Hybrid", stats['cobol_hybrid'])
            
            st.markdown("---")
            st.markdown("**🗄️ Ressources**")
            
            col11, col12, col13, col14 = st.columns(4)
            col11.metric("🎫 Transactions", stats['total_transactions'])
            col12.metric("🗄️ Tables DB2", stats['total_db2_tables'])
            col13.metric("📊 Segments IMS", stats['total_ims_segments'])
            col14.metric("📨 Queues MQ", stats['total_mq_queues'])
            
            st.markdown("---")
            st.markdown("**⚠️ Indicateurs de risque**")
            
            col15, col16, col17, col18 = st.columns(4)
            col15.metric("🚫 Orphelins", stats['orphan_programs'], delta_color="inverse")
            col16.metric("⚠️ Critiques", stats['critical_programs'], delta_color="inverse")
            col17.metric("🔥 Haute complexité", stats['high_risk_programs'], delta_color="inverse")
            col18.metric("🔄 Cycles", len(cycles), delta_color="inverse")
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Graphiques
            try:
                import matplotlib.pyplot as plt
                
                st.markdown('<div class="glass-card">', unsafe_allow_html=True)
                st.subheader("📈 Visualisations")
                
                fig, axes = plt.subplots(1, 2, figsize=(14, 5))
                fig.patch.set_facecolor('#1E1E2E')
                
                # Graph 1
                types = ['Batch', 'CICS', 'IMS', 'DB2', 'Hybrid']
                counts = [stats['cobol_batch'], stats['cobol_cics'], stats['cobol_ims'], stats['cobol_db2'], stats['cobol_hybrid']]
                colors = ['#4EA3FF', '#00D9A5', '#FFB020', '#FF6B9D', '#A3FFD6']
                
                axes[0].pie(counts, labels=types, autopct='%1.1f%%', colors=colors, textprops={'color': 'white'})
                axes[0].set_title('Répartition par Type', color='white', fontsize=12, pad=20)
                axes[0].set_facecolor('#262636')
                
                # Graph 2
                if analyzer.programs:
                    top_complex = sorted(analyzer.programs.values(), key=lambda p: p.complexity_score, reverse=True)[:10]
                    names = [p.name[:12] for p in top_complex]
                    scores = [p.complexity_score for p in top_complex]
                    
                    axes[1].barh(names, scores, color='#FF6B9D')
                    axes[1].set_title('Top 10 Complexité', color='white', fontsize=12, pad=20)
                    axes[1].set_xlabel('Score', color='white')
                    axes[1].tick_params(colors='white')
                    axes[1].set_facecolor('#262636')
                
                plt.tight_layout()
                st.pyplot(fig)
                
                st.markdown('</div>', unsafe_allow_html=True)
            except:
                pass

        # TAB 2 : PROGRAMMES
        with tab2:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("📘 Inventaire Programmes")
            
            if analyzer.programs:
                df_programs = pd.DataFrame([
                    {
                        'Nom': p.name,
                        'Type': p.component_type.value.replace('COBOL_', ''),
                        'Lignes': p.lines,
                        'Paragraphes': p.paragraphs,
                        'Complexité': p.complexity_score,
                        'Risque': p.risk_level,
                        'Appelé par': len(p.called_by),
                        'Appelle': len(p.calls_to),
                        'Orphelin': '🚫' if p.is_orphan else '',
                        'Critique': '⚠️' if p.is_critical else '',
                        'CICS': '✓' if p.calls_cics else '',
                        'DB2': len(p.db2_tables) if p.db2_tables else '',
                        'IA': '🤖' if p.name in llm_analyses else ''
                    }
                    for p in analyzer.programs.values()
                ])
                
                st.dataframe(df_programs.sort_values('Complexité', ascending=False), use_container_width=True, height=500)
            
            st.markdown('</div>', unsafe_allow_html=True)

        # TAB 3 : ANALYSES IA (NOUVEAU)
        with tab3:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("🤖 Analyses Intelligentes par IA")
            
            if llm_analyses:
                st.markdown(f"""
                <div class="success-box">
                    ✅ <strong>{len(llm_analyses)} programme(s) analysé(s) par l'IA</strong>
                </div>
                """, unsafe_allow_html=True)
                
                # Sélecteur de programme
                selected_prog = st.selectbox(
                    "Sélectionnez un programme pour voir l'analyse détaillée :",
                    options=sorted(llm_analyses.keys()),
                    key="ultimate_llm_prog_select"
                )
                
                if selected_prog:
                    st.markdown("---")
                    st.markdown(f"### 📋 Analyse de : **{selected_prog}**")
                    
                    # Afficher l'analyse LLM
                    analysis_text = llm_analyses[selected_prog]
                    
                    # Parser l'analyse pour affichage structuré
                    if "RÔLE MÉTIER:" in analysis_text:
                        sections = {
                            'RÔLE MÉTIER:': '💼',
                            'FLUX DE DONNÉES:': '📊',
                            'DÉPENDANCES CRITIQUES:': '🔗',
                            'NIVEAU DE RISQUE:': '⚠️',
                            'RECOMMANDATIONS:': '💡'
                        }
                        
                        for section_name, icon in sections.items():
                            if section_name in analysis_text:
                                st.markdown(f"**{icon} {section_name.replace(':', '')}**")
                                
                                # Extraire le contenu de cette section
                                start_idx = analysis_text.find(section_name) + len(section_name)
                                
                                # Trouver la prochaine section
                                next_sections = [analysis_text.find(s) for s in sections.keys() if s != section_name and analysis_text.find(s) > start_idx]
                                end_idx = min(next_sections) if next_sections else len(analysis_text)
                                
                                content = analysis_text[start_idx:end_idx].strip()
                                
                                # Afficher dans une box selon le type
                                if 'RISQUE' in section_name:
                                    if 'ÉLEVÉ' in content.upper():
                                        st.markdown(f'<div class="error-box">{content}</div>', unsafe_allow_html=True)
                                    elif 'MOYEN' in content.upper():
                                        st.markdown(f'<div class="warning-box">{content}</div>', unsafe_allow_html=True)
                                    else:
                                        st.markdown(f'<div class="success-box">{content}</div>', unsafe_allow_html=True)
                                else:
                                    st.markdown(f'<div class="info-box">{content}</div>', unsafe_allow_html=True)
                                
                                st.markdown("---")
                    else:
                        # Affichage brut si parsing échoue
                        st.markdown(analysis_text)
                    
                    # Infos techniques du programme
                    if selected_prog in analyzer.programs:
                        prog = analyzer.programs[selected_prog]
                        
                        with st.expander("🔧 Détails techniques"):
                            col_tech1, col_tech2, col_tech3 = st.columns(3)
                            
                            with col_tech1:
                                st.markdown("**Métriques**")
                                st.text(f"Lignes: {prog.lines}")
                                st.text(f"Paragraphes: {prog.paragraphs}")
                                st.text(f"Complexité: {prog.complexity_score}")
                            
                            with col_tech2:
                                st.markdown("**Appels**")
                                st.text(f"CALL: {len(prog.calls_static)}")
                                st.text(f"CICS: {len(prog.calls_cics)}")
                                st.text(f"DB2: {len(prog.db2_tables)}")
                            
                            with col_tech3:
                                st.markdown("**Ressources**")
                                st.text(f"Copybooks: {len(prog.copybooks)}")
                                st.text(f"Trans CICS: {len(prog.transactions)}")
                                st.text(f"IMS Seg: {len(prog.ims_segments)}")
            else:
                st.info("Aucune analyse IA disponible. Activez l'option lors de l'analyse.")
            
            st.markdown('</div>', unsafe_allow_html=True)

        # TAB 4 : GRAPHE
        with tab4:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("🔗 Graphe de Dépendances")
            
            if graph:
                try:
                    import networkx as nx
                    import matplotlib.pyplot as plt
                    
                    fig, ax = plt.subplots(figsize=(16, 12))
                    fig.patch.set_facecolor('#1E1E2E')
                    ax.set_facecolor('#262636')
                    
                    pos = nx.spring_layout(graph, k=3, iterations=50, seed=42)
                    
                    node_colors = []
                    for node in graph.nodes():
                        node_data = graph.nodes[node]
                        node_type = node_data.get('type', 'PROGRAM')
                        
                        if node_type == 'DB2_TABLE':
                            node_colors.append('#FFB020')
                        elif node_type == 'IMS_SEGMENT':
                            node_colors.append('#FF6B9D')
                        elif node_type == 'MQ_QUEUE':
                            node_colors.append('#A3FFD6')
                        else:
                            is_orphan = node_data.get('is_orphan', False)
                            is_critical = node_data.get('is_critical', False)
                            
                            if is_orphan:
                                node_colors.append('#FF4757')
                            elif is_critical:
                                node_colors.append('#FF6B00')
                            else:
                                node_colors.append('#4EA3FF')
                    
                    nx.draw_networkx_nodes(graph, pos, node_color=node_colors, node_size=1000, ax=ax, alpha=0.9)
                    nx.draw_networkx_labels(graph, pos, font_size=7, font_color='white', ax=ax, font_weight='bold')
                    nx.draw_networkx_edges(graph, pos, edge_color='#00D9A5', arrows=True, arrowsize=20, width=2, ax=ax, alpha=0.6)
                    
                    ax.set_title('Graphe de Dépendances', color='white', fontsize=16, pad=20)
                    ax.axis('off')
                    
                    plt.tight_layout()
                    st.pyplot(fig)
                    
                    # Statistiques
                    col_graph1, col_graph2, col_graph3 = st.columns(3)
                    col_graph1.metric("Nœuds", graph.number_of_nodes())
                    col_graph2.metric("Arêtes", graph.number_of_edges())
                    col_graph3.metric("Cycles", len(cycles), delta_color="inverse")
                
                except:
                    st.error("❌ Erreur affichage graphe")
            else:
                st.info("Graphe non généré")
            
            st.markdown('</div>', unsafe_allow_html=True)

        # TAB 5 : IMPACT
        with tab5:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("🔍 Recherche d'Impact")
            
            program_to_analyze = st.selectbox(
                "Programme à modifier :",
                options=sorted(analyzer.programs.keys()),
                key="ultimate_impact_select"
            )
            
            if st.button("🔍 Analyser", key="ultimate_impact_btn"):
                impact_result = analyzer.search_impact(program_to_analyze)
                
                if 'error' not in impact_result:
                    st.markdown(f"""
                    <div class="info-box">
                        📊 <strong>Programmes impactés :</strong> {impact_result['impact_count']}<br>
                        ⚠️ <strong>Risque :</strong> {impact_result['risk_assessment']}
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if impact_result['direct_impact']:
                        st.markdown("**Impact direct :**")
                        for prog in impact_result['direct_impact']:
                            st.markdown(f"- 🔴 {prog}")
            
            st.markdown('</div>', unsafe_allow_html=True)

        # TAB 6 : ALERTES
        with tab6:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("⚠️ Alertes et Risques")
            
            orphans = [p for p in analyzer.programs.values() if p.is_orphan]
            criticals = [p for p in analyzer.programs.values() if p.is_critical]
            high_complex = [p for p in analyzer.programs.values() if p.risk_level == 'HIGH']
            
            if orphans:
                st.markdown(f"""
                <div class="warning-box">
                    🚫 <strong>{len(orphans)} orphelin(s)</strong>
                </div>
                """, unsafe_allow_html=True)
            
            if criticals:
                st.markdown(f"""
                <div class="error-box">
                    ⚠️ <strong>{len(criticals)} critique(s)</strong>
                </div>
                """, unsafe_allow_html=True)
            
            if high_complex:
                st.markdown(f"""
                <div class="error-box">
                    🔥 <strong>{len(high_complex)} haute complexité</strong>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)

        # TAB 7 : MIGRATION (NOUVEAU)
        with tab7:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("🚀 Plan de Migration")
            
            if migration_plan:
                st.markdown(migration_plan)
            else:
                st.info("Plan de migration non généré")
            
            st.markdown('</div>', unsafe_allow_html=True)

        # TAB 8 : EXPORTS
        with tab8:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.subheader("💾 Exports")
            
            col_exp1, col_exp2, col_exp3 = st.columns(3)
            
            # JSON
            with col_exp1:
                json_data = json.dumps(report_json, indent=2, ensure_ascii=False)
                st.download_button(
                    "📄 JSON",
                    data=json_data.encode('utf-8'),
                    file_name=f"analysis_{datetime.now().strftime('%Y%m%d')}.json",
                    key="ultimate_download_json"
                )
            
            # Excel
            with col_exp2:
                try:
                    excel_buf = io.BytesIO()
                    with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
                        if not df_programs.empty:
                            df_programs.to_excel(writer, index=False, sheet_name="Programmes")
                        
                        df_stats = pd.DataFrame(list(stats.items()), columns=['Indicateur', 'Valeur'])
                        df_stats.to_excel(writer, index=False, sheet_name="Stats")
                    
                    excel_buf.seek(0)
                    
                    st.download_button(
                        "📊 Excel",
                        data=excel_buf,
                        file_name=f"analysis_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="ultimate_download_excel"
                    )
                except:
                    st.info("pip install xlsxwriter")
            
            # ZIP
            with col_exp3:
                try:
                    zip_buf = io.BytesIO()
                    with zipfile.ZipFile(zip_buf, 'w') as zipf:
                        zipf.writestr("report.json", json_data)
                        
                        if migration_plan:
                            zipf.writestr("migration_plan.md", migration_plan)
                    
                    zip_buf.seek(0)
                    
                    st.download_button(
                        "📦 ZIP",
                        data=zip_buf,
                        file_name=f"analysis_{datetime.now().strftime('%Y%m%d')}.zip",
                        mime="application/zip",
                        key="ultimate_download_zip"
                    )
                except:
                    pass
            
            st.markdown('</div>', unsafe_allow_html=True)
# ===================== FOOTER PRO =====================
st.markdown("""
<div class="footer-pro">
    <div class="footer-title">💼 Modernisation Mainframe IBM A&T</div>
    <div class="footer-team">
        👥 <span>Équipe</span> : Younes • Hanane • Nezha • Aimane • Khaoula • Naoufal • Imane • Meriem
    </div>
    <div style="margin-top: 1rem; color: #666; font-size: 0.85rem;">
         © 2025 Tous droits réservés
    </div>
</div>
""", unsafe_allow_html=True)

# ---- Hide Streamlit Footer ----
hide_footer_style = """
    <style>
    footer {visibility: hidden;}      /* Cache le footer */
    </style>
"""
st.markdown(hide_footer_style, unsafe_allow_html=True)