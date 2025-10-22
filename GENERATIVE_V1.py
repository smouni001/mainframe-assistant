# GENAIV1_CLAUDE.py — Mainframe Modernization AI Assistant (PRO VERSION)
# ========================================================================

import io
import os
import zipfile
from datetime import datetime
from functools import lru_cache
from typing import Optional, List, Tuple

import pandas as pd
import streamlit as st

# ==== LLM client ====
try:
    from langchain_anthropic import ChatAnthropic
    CLAUDE_AVAILABLE = True
except ImportError:
    ChatAnthropic = None
    CLAUDE_AVAILABLE = False

# ==== Langchain utils ====
try:
    from langchain_openai import OpenAIEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document as LCDocument
    from langchain_core.prompts import PromptTemplate
    LANGCHAIN_AVAILABLE = True
except ImportError:
    OpenAIEmbeddings = None
    FAISS = None
    LANGCHAIN_AVAILABLE = False

# ===================== CONFIGURATION =====================
st.set_page_config(
    page_title="Assistant AI Mainframe Pro",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ===================== CUSTOM CSS PRO =====================
st.markdown("""
<style>
    /* ===== THEME PRINCIPAL ===== */
    :root {
        --primary-color: #4EA3FF;
        --secondary-color: #0066CC;
        --accent-color: #00D4FF;
        --bg-dark: #0E1117;
        --bg-card: #1E1E2E;
        --text-primary: #FFFFFF;
        --text-secondary: #B8B8C8;
        --success: #00D9A5;
        --warning: #FFB020;
        --error: #FF4757;
    }
    
    /* ===== BACKGROUND ===== */
    .stApp {
        background: linear-gradient(135deg, #0E1117 0%, #1a1d29 100%);
    }
    
    /* ===== HEADER PRO ===== */
    .ai-header {
        background: linear-gradient(135deg, #1E1E2E 0%, #2B2D3E 100%);
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 8px 32px rgba(78, 163, 255, 0.15);
        border: 1px solid rgba(78, 163, 255, 0.2);
        position: relative;
        overflow: hidden;
    }
    
    .ai-header::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(78, 163, 255, 0.1), transparent);
        animation: shine 3s infinite;
    }
    
    @keyframes shine {
        0% { left: -100%; }
        100% { left: 100%; }
    }
    
    .ai-header h1 {
        font-size: 2.5rem;
        font-weight: 800;
        margin: 0;
        background: linear-gradient(135deg, #4EA3FF 0%, #00D4FF 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        position: relative;
        z-index: 1;
    }
    
    .ai-header .subtitle {
        color: var(--text-secondary);
        font-size: 1rem;
        margin-top: 0.5rem;
        font-weight: 300;
        letter-spacing: 1px;
    }
    
    /* ===== CARDS GLASSMORPHISM ===== */
    .glass-card {
        background: rgba(30, 30, 46, 0.7);
        backdrop-filter: blur(10px);
        border-radius: 12px;
        padding: 1.5rem;
        border: 1px solid rgba(78, 163, 255, 0.15);
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
        margin-bottom: 1.5rem;
        transition: all 0.3s ease;
    }
    
    .glass-card:hover {
        border-color: rgba(78, 163, 255, 0.4);
        box-shadow: 0 12px 40px rgba(78, 163, 255, 0.2);
        transform: translateY(-2px);
    }
    
    /* ===== INFO BOXES ===== */
    .info-box {
        background: linear-gradient(135deg, rgba(78, 163, 255, 0.1) 0%, rgba(0, 212, 255, 0.1) 100%);
        border-left: 4px solid var(--primary-color);
        padding: 1rem 1.5rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: var(--text-primary);
        font-size: 0.95rem;
        line-height: 1.6;
    }
    
    .success-box {
        background: linear-gradient(135deg, rgba(0, 217, 165, 0.1) 0%, rgba(0, 255, 200, 0.1) 100%);
        border-left-color: var(--success);
    }
    
    .warning-box {
        background: linear-gradient(135deg, rgba(255, 176, 32, 0.1) 0%, rgba(255, 200, 100, 0.1) 100%);
        border-left-color: var(--warning);
    }
    
    .error-box {
        background: linear-gradient(135deg, rgba(255, 71, 87, 0.1) 0%, rgba(255, 100, 120, 0.1) 100%);
        border-left-color: var(--error);
    }
    
    /* ===== BUTTONS PRO ===== */
    .stButton>button {
        background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%);
        color: white;
        font-weight: 600;
        border-radius: 10px;
        border: none;
        padding: 0.75rem 2rem;
        font-size: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(78, 163, 255, 0.3);
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, var(--secondary-color) 0%, var(--primary-color) 100%);
        box-shadow: 0 6px 25px rgba(78, 163, 255, 0.5);
        transform: translateY(-2px);
    }
    
    .stButton>button:active {
        transform: translateY(0);
    }
    
    .stButton>button:disabled {
        background: linear-gradient(135deg, #2a2a3a 0%, #3a3a4a 100%);
        box-shadow: none;
        opacity: 0.5;
    }
    
    /* ===== SIDEBAR PRO ===== */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1a1d29 0%, #0E1117 100%);
        border-right: 1px solid rgba(78, 163, 255, 0.15);
    }
    
    [data-testid="stSidebar"] .stRadio > label {
        background: rgba(78, 163, 255, 0.05);
        padding: 0.5rem 1rem;
        border-radius: 8px;
        margin: 0.25rem 0;
        transition: all 0.2s ease;
    }
    
    [data-testid="stSidebar"] .stRadio > label:hover {
        background: rgba(78, 163, 255, 0.15);
    }
    
    /* ===== INPUTS ===== */
    .stTextInput input, .stTextArea textarea {
        background: rgba(30, 30, 46, 0.6) !important;
        border: 1px solid rgba(78, 163, 255, 0.3) !important;
        border-radius: 8px !important;
        color: var(--text-primary) !important;
        padding: 0.75rem !important;
        transition: all 0.3s ease !important;
    }
    
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--primary-color) !important;
        box-shadow: 0 0 0 2px rgba(78, 163, 255, 0.2) !important;
    }
    
    /* ===== FILE UPLOADER ===== */
    [data-testid="stFileUploader"] {
        background: rgba(30, 30, 46, 0.4);
        border: 2px dashed rgba(78, 163, 255, 0.3);
        border-radius: 12px;
        padding: 2rem;
        transition: all 0.3s ease;
    }
    
    [data-testid="stFileUploader"]:hover {
        border-color: var(--primary-color);
        background: rgba(78, 163, 255, 0.05);
    }
    
    /* ===== DATAFRAME ===== */
    .stDataFrame {
        border-radius: 12px;
        overflow: hidden;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
    }
    
    /* ===== CODE BLOCKS ===== */
    .stCodeBlock {
        border-radius: 12px;
        border: 1px solid rgba(78, 163, 255, 0.2);
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
    }
    
    /* ===== SPINNER ===== */
    .stSpinner > div {
        border-top-color: var(--primary-color) !important;
    }
    
    /* ===== TABS ===== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: rgba(30, 30, 46, 0.5);
        padding: 0.5rem;
        border-radius: 12px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: transparent;
        border-radius: 8px;
        color: var(--text-secondary);
        padding: 0.75rem 1.5rem;
        transition: all 0.3s ease;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%);
        color: white;
    }
    
    /* ===== FOOTER PRO ===== */
    .footer-pro {
        margin-top: 4rem;
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        background: linear-gradient(135deg, #1E1E2E 0%, #2B2D3E 100%);
        border: 1px solid rgba(78, 163, 255, 0.2);
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
    }
    
    .footer-title {
        color: var(--primary-color);
        font-weight: 700;
        font-size: 1.1rem;
        margin-bottom: 1rem;
    }
    
    .footer-team {
        color: var(--text-secondary);
        font-size: 0.9rem;
        line-height: 1.8;
    }
    
    .footer-team span {
        color: var(--accent-color);
        font-weight: 600;
    }
    
    /* ===== ANIMATIONS ===== */
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .glass-card, .info-box {
        animation: fadeIn 0.5s ease-out;
    }
    
    /* ===== RESPONSIVE ===== */
    @media (max-width: 768px) {
        .ai-header h1 {
            font-size: 1.8rem;
        }
        
        .glass-card {
            padding: 1rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# ===================== HEADER PRO =====================
st.markdown("""
<div class="ai-header">
    <h1>🤖 Assistant AI Mainframe</h1>
    <div class="subtitle">IBM A&T</div>
</div>
""", unsafe_allow_html=True)

# ===================== LANGUE =====================
lang_choice = st.sidebar.selectbox(
    "🌐 Language / Langue", 
    ["Français", "English"], 
    index=0,
    help="Sélectionnez votre langue préférée"
)
LANG_FR = (lang_choice == "Français")

def T(fr: str, en: str) -> str:
    return fr if LANG_FR else en

# ===================== API KEY =====================

ANTHROPIC_API_KEY = st.secrets["ANTHROPIC_API_KEY"]


if not CLAUDE_AVAILABLE or not ANTHROPIC_API_KEY:
    st.markdown("""
    <div class="error-box">
        ⚠️ <strong>LLM Claude indisponible</strong><br>
        Installez <code>langchain-anthropic</code> et configurez votre clé API
    </div>
    """, unsafe_allow_html=True)
    st.stop()

# ===================== PROMPT ENGINE =====================
import yaml

@lru_cache(maxsize=1)
def load_prompt_engine(path: str = "PromptEngine.yaml") -> dict:
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
        else:
            st.markdown(f"""
            <div class="warning-box">
                ℹ️ Fichier <code>{path}</code> introuvable. Utilisation des prompts par défaut.
            </div>
            """, unsafe_allow_html=True)
            return {}
    except Exception as e:
        st.markdown(f"""
        <div class="error-box">
            ❌ Erreur lors du chargement de <code>{path}</code>: {e}
        </div>
        """, unsafe_allow_html=True)
        return {}

def get_prompt(section: str, lang_key: Optional[str] = None, **kwargs) -> str:
    pe = load_prompt_engine()
    defaults = pe.get("defaults", {})
    invalid_message = defaults.get("invalid_message", "❌ Code invalide")
    no_jcl = defaults.get("no_jcl_in_test_mode", "Ne pas générer de JCL")

    if section == "ANALYSE_DOC":
        node = pe.get("ANALYSE_DOC", {}).get("GENERIQUE", "")
    elif lang_key:
        node = pe.get(section, {}).get(lang_key.upper(), "")
    else:
        node = ""

    if not node:
        if section == "ANALYSE_DOC":
            node = "Rôle: Analyste documentaire.\n{document_content}"
        elif section == "JCL_GENERATION":
            node = "Rôle: Expert JCL {lang}.\nCode:\n{source_code}"
        else:
            node = "Rôle: QA Test {lang}.\nModule:{module_base}\nCode:\n{source_code}"

    try:
        return node.format(
            invalid_message=invalid_message,
            no_jcl_in_test_mode=no_jcl,
            **kwargs
        )
    except KeyError as e:
        st.error(f"❌ Paramètre manquant: {e}")
        return node

# ===================== HELPERS =====================
def llm_client(max_tokens: int = 2200, temperature: float = 0.2):
    if not CLAUDE_AVAILABLE or not ANTHROPIC_API_KEY:
        return None
    try:
        return ChatAnthropic(
            api_key=ANTHROPIC_API_KEY,
            temperature=temperature,
            max_tokens=max_tokens,
            model_name="claude-3-haiku-20240307"
        )
    except Exception as e:
        st.error(f"❌ Erreur client Claude: {e}")
        return None

def _normalize_text(s: str) -> str:
    return " ".join(s.replace("\r", "\n").split())

def read_pdf_bytes(data: bytes) -> str:
    try:
        import PyPDF2
        text = []
        with io.BytesIO(data) as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text.append(extracted)
        result = _normalize_text("\n".join(text))
        if result.strip():
            return result
    except:
        pass
    
    try:
        import pdfplumber
        text = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text.append(extracted)
        return _normalize_text("\n".join(text))
    except:
        return ""

def read_docx_bytes(data: bytes) -> str:
    try:
        import docx
        with io.BytesIO(data) as fh:
            doc = docx.Document(fh)
            return _normalize_text("\n".join(p.text for p in doc.paragraphs))
    except:
        return ""

def read_txt_bytes(data: bytes, encoding: str = "utf-8") -> str:
    try:
        return _normalize_text(data.decode(encoding, errors="ignore"))
    except:
        return ""

def extract_document_content(uploaded_file) -> str:
    if not uploaded_file:
        return ""
    
    name = (uploaded_file.name or "").lower()
    
    try:
        raw = uploaded_file.read()
    except Exception as e:
        st.error(f"❌ Erreur de lecture: {e}")
        return ""

    def safe_truncate(text: str, max_chars: int = 150_000) -> str:
        if len(text) > max_chars:
            st.markdown(f"""
            <div class="warning-box">
                ⚠️ Contenu tronqué à {max_chars:,} caractères
            </div>
            """, unsafe_allow_html=True)
        return text[:max_chars]

    if name.endswith(".pdf"):
        return safe_truncate(read_pdf_bytes(raw))
    if name.endswith(".docx"):
        return safe_truncate(read_docx_bytes(raw))
    if name.endswith(".txt"):
        return safe_truncate(read_txt_bytes(raw))

    if name.endswith(".zip"):
        texts = []
        try:
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                for info in z.infolist():
                    n = info.filename.lower()
                    if n.endswith((".pdf", ".docx", ".txt")) and not info.is_dir():
                        try:
                            data = z.read(info)
                            if n.endswith(".pdf"):
                                texts.append(read_pdf_bytes(data))
                            elif n.endswith(".docx"):
                                texts.append(read_docx_bytes(data))
                            elif n.endswith(".txt"):
                                texts.append(read_txt_bytes(data))
                            
                            if sum(len(t) for t in texts) > 200_000:
                                break
                        except Exception as e:
                            st.warning(f"⚠️ Impossible de lire {info.filename}")
                            continue
        except zipfile.BadZipFile:
            st.markdown("""
            <div class="error-box">
                ❌ Fichier ZIP invalide ou corrompu
            </div>
            """, unsafe_allow_html=True)
            return ""
        
        return safe_truncate("\n\n---\n\n".join(t for t in texts if t.strip()))

    return ""

# ===================== UI LABELS =====================
TEXTS = {
    "Français": {
        "choose_mode": "⚙️ Mode de traitement",
        "modes": ["📄 Analyse documentaire", "🔧 Génération JCL", "🧪 Test COBOL"],
    },
    "English": {
        "choose_mode": "⚙️ Processing Mode",
        "modes": ["📄 Document Analysis", "🔧 JCL Generation", "🧪 COBOL Testing"],
    }
}
TXT = TEXTS["Français"] if LANG_FR else TEXTS["English"]

# ===================== MODE SELECTOR =====================
st.sidebar.markdown("---")
mode = st.sidebar.radio(
    TXT["choose_mode"], 
    TXT["modes"], 
    index=0,
    help=T("Sélectionnez le mode de traitement souhaité", "Select your processing mode")
)

# ===================== MODE 1 : ANALYSE DOC =====================
if mode == TXT["modes"][0]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header(T("📄 Analyse Documentaire Intelligente", "📄 Intelligent Document Analysis"))
    st.markdown('</div>', unsafe_allow_html=True)

    # Upload section
    col1, col2 = st.columns(2)
    with col1:
        uploaded_zip = st.file_uploader(
            "📦 " + T("Fichier ZIP", "ZIP File"),
            type=["zip"],
            help=T("Charger un fichier ZIP contenant des documents", "Upload a ZIP file with documents")
        )
    with col2:
        uploaded_pdfs = st.file_uploader(
            "📄 " + T("Fichiers PDF", "PDF Files"),
            type=["pdf"],
            accept_multiple_files=True,
            help=T("Charger un ou plusieurs PDFs", "Upload one or more PDFs")
        )

    pdf_buffers: List[Tuple[str, bytes]] = []
    
    if uploaded_zip or uploaded_pdfs:
        with st.spinner(T("📚 Traitement des documents...", "📚 Processing documents...")):
            if uploaded_zip:
                try:
                    with zipfile.ZipFile(uploaded_zip, "r") as zip_ref:
                        for name in zip_ref.namelist():
                            if name.lower().endswith(".pdf") and not name.startswith("__MACOSX"):
                                try:
                                    with zip_ref.open(name) as pdf_file:
                                        pdf_buffers.append((name, pdf_file.read()))
                                except Exception as e:
                                    st.warning(f"⚠️ {name}: {e}")
                except zipfile.BadZipFile:
                    st.markdown('<div class="error-box">❌ ZIP invalide</div>', unsafe_allow_html=True)
            
            if uploaded_pdfs:
                for f in uploaded_pdfs:
                    try:
                        pdf_buffers.append((f.name, f.read()))
                    except Exception as e:
                        st.warning(f"⚠️ {f.name}: {e}")
    
    if pdf_buffers:
        st.markdown(f"""
        <div class="success-box">
            ✅ <strong>{len(pdf_buffers)} document(s) chargé(s)</strong>
        </div>
        """, unsafe_allow_html=True)

    # Analysis context
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    context_mode = st.radio(
        "🎯 " + T("Contexte d'analyse", "Analysis Context"),
        [T("Analyse Générique", "Generic Analysis"),
         T("Analyse CV / Profil", "CV / Profile Analysis")],
        horizontal=True
    )

    question = st.text_area(
        "💬 " + T("Votre question", "Your question"),
        placeholder=T(
            "Ex: Quelles sont les compétences principales ?",
            "Ex: What are the main skills?"
        ),
        height=100
    )
    st.markdown('</div>', unsafe_allow_html=True)

    if st.button(
        "🚀 " + T("ANALYSER", "ANALYZE"), 
        disabled=not question or not pdf_buffers,
        use_container_width=True
    ):
        if not pdf_buffers:
            st.markdown('<div class="error-box">❌ Aucun document chargé</div>', unsafe_allow_html=True)
        else:
            texts = []
            for name, raw in pdf_buffers:
                extracted = read_pdf_bytes(raw)
                if extracted.strip():
                    texts.append(f"=== {name} ===\n{extracted}")
            
            document_content = "\n\n---\n\n".join(texts)[:150_000]

            if not document_content.strip():
                st.markdown('<div class="error-box">❌ Aucun texte extrait</div>', unsafe_allow_html=True)
            else:
                if "CV" in context_mode or "Profile" in context_mode:
                    prompt_text = f"""
Tu es une IA experte en analyse de CV et matching de profils.

**Mission :**
- Analyse approfondie des documents fournis
- Identification des compétences (explicites et implicites)
- Évaluation complète avant réponse
- Aucun détail important omis

**Format de réponse :**
1. Réponse directe
2. Détails et justifications
3. Score (si applicable) : X/100
4. Recommandations

**Question :** {question}

**Documents :**
{document_content}
"""
                else:
                    prompt_text = get_prompt(
                        "ANALYSE_DOC",
                        None,
                        document_content=document_content,
                        question=question
                    )

                client = llm_client(max_tokens=3000, temperature=0.2)
                if not client:
                    st.markdown('<div class="error-box">❌ Client LLM indisponible</div>', unsafe_allow_html=True)
                else:
                    with st.spinner(T("🧠 Analyse en cours...", "🧠 Analyzing...")):
                        try:
                            response = client.invoke(prompt_text)
                            result = response.content if hasattr(response, 'content') else str(response)
                        except Exception as e:
                            st.error(f"❌ Erreur LLM: {e}")
                            result = None

                    if result:
                        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
                        st.subheader("🧠 " + T("Réponse de l'IA", "AI Answer"))
                        st.markdown(result)
                        st.markdown('</div>', unsafe_allow_html=True)

                        # Export
                        try:
                            from docx import Document
                            doc = Document()
                            doc.add_heading(T("Rapport d'Analyse IA", "AI Analysis Report"), 0)
                            doc.add_heading(T("Question", "Question"), 1)
                            doc.add_paragraph(question)
                            doc.add_heading(T("Réponse", "Answer"), 1)
                            doc.add_paragraph(result)
                            
                            buf = io.BytesIO()
                            doc.save(buf)
                            buf.seek(0)
                            
                            st.download_button(
                                "📥 " + T("Télécharger Rapport (Word)", "Download Report (Word)"),
                                data=buf,
                                file_name=T("rapport_analyse.docx", "analysis_report.docx"),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        except ImportError:
                            st.download_button(
                                "📥 " + T("Télécharger Rapport (TXT)", "Download Report (TXT)"),
                                data=result.encode("utf-8"),
                                file_name=T("rapport_analyse.txt", "analysis_report.txt"),
                                use_container_width=True
                            )

# ===================== MODE 2 : JCL GENERATION =====================
elif mode == TXT["modes"][1]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header("🔧 " + T("Générateur JCL Automatique", "Automatic JCL Generator"))
    st.markdown('</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "📂 " + T("Fichier programme", "Program file"),
        type=["cbl", "cob", "pli", "pl1", "asm"],
        help=T("Sélectionnez votre fichier COBOL/PL/I/ASM", "Select your COBOL/PL/I/ASM file")
    )
    
    if uploaded_file:
        st.markdown(f"""
        <div class="success-box">
            ✅ Fichier chargé : <strong>{uploaded_file.name}</strong>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("⚙️ " + T("Paramètres z/OS", "z/OS Parameters"))
    
    col1, col2, col3 = st.columns(3)
    with col1:
        job_name = st.text_input("🧾 JOB", value="GENJOB1")
        program_name = st.text_input("🏷️ " + T("Programme", "Program"), value="PROGTEST")
    with col2:
        pds_source = st.text_input("📁 PDS Source", value="MYUSER.COBOL.SOURCE")
        loadlib = st.text_input("💾 LOADLIB", value="MYUSER.COBOL.LOAD")
    with col3:
        sysout_class = st.text_input("🖨️ SYSOUT", value="A")
        region_size = st.text_input("💽 REGION", value="4096K")
    
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown(f"""
    <div class="info-box">
        📊 <strong>Configuration :</strong> JOB={job_name} | PGM={program_name} | PDS={pds_source}
    </div>
    """, unsafe_allow_html=True)

    if st.button(
        "🚀 " + T("GÉNÉRER JCL", "GENERATE JCL"),
        disabled=not uploaded_file,
        use_container_width=True
    ):
        file_name = uploaded_file.name
        ext = file_name.split(".")[-1].lower()
        
        try:
            source_code = uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception as e:
            st.error(f"❌ Erreur: {e}")
            st.stop()

        lang_map = {"cbl": "COBOL", "cob": "COBOL", "pli": "PL/I", "pl1": "PL/I", "asm": "ASM"}
        lang = lang_map.get(ext)

        if not lang:
            st.markdown('<div class="error-box">⚠️ Type non supporté</div>', unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div class="info-box">
                🔍 Type détecté : <strong>{lang}</strong>
            </div>
            """, unsafe_allow_html=True)
            
            prompt_text = get_prompt(
                "JCL_GENERATION", lang,
                source_code=source_code,
                job_name=job_name,
                program_name=program_name,
                pds_source=pds_source,
                loadlib=loadlib,
                sysout_class=sysout_class,
                region_size=region_size
            )
            
            client = llm_client(max_tokens=2500, temperature=0.3)
            if client:
                with st.spinner(T("🧠 Génération du JCL...", "🧠 Generating JCL...")):
                    try:
                        response = client.invoke(prompt_text)
                        result = response.content if hasattr(response, 'content') else str(response)
                        
                        if result:
                            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
                            st.subheader("📘 " + T("JCL Généré", "Generated JCL"))
                            st.code(result, language="jcl")
                            st.markdown('</div>', unsafe_allow_html=True)
                            
                            st.download_button(
                                "💾 " + T("Télécharger JCL", "Download JCL"),
                                data=result.encode("utf-8"),
                                file_name=f"job_{lang.lower()}.jcl",
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"❌ Erreur: {e}")

# ===================== MODE 3 : COBOL TEST =====================
else:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.header("🧪 " + T("Générateur de Tests COBOL", "COBOL Test Generator"))
    st.markdown('</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "📂 " + T("Module source", "Source module"),
        type=["cbl", "cob", "pli", "pl1", "asm"],
        help=T("Sélectionnez votre module", "Select your module")
    )
    
    if uploaded_file:
        st.markdown(f"""
        <div class="success-box">
            ✅ Module chargé : <strong>{uploaded_file.name}</strong>
        </div>
        """, unsafe_allow_html=True)

        file_name = uploaded_file.name
        module_base = file_name.rsplit('.', 1)[0]
        ext = file_name.split(".")[-1].lower()
        
        try:
            source_code = uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception as e:
            st.error(f"❌ Erreur: {e}")
            st.stop()

        lang_map = {"cbl": "COBOL", "cob": "COBOL", "pli": "PL/I", "pl1": "PL/I", "asm": "ASM"}
        language_type = lang_map.get(ext)

        if not language_type:
            st.markdown('<div class="error-box">⚠️ Type non supporté</div>', unsafe_allow_html=True)
            st.stop()

        module_starts_ok = module_base.upper().startswith("M")
        if not module_starts_ok:
            st.markdown(f"""
            <div class="warning-box">
                ℹ️ Le nom du module doit commencer par <strong>M</strong><br>
                Exemple : <code>MTRAITSIM.cbl</code>
            </div>
            """, unsafe_allow_html=True)

        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        default_test_name = (module_base.upper() + "TEST")
        if not default_test_name.startswith("M"):
            default_test_name = "M" + default_test_name
        default_test_name = "".join(ch for ch in default_test_name if ch.isalnum())[:8]

        test_prog_name = st.text_input(
            "🏷️ " + T("Nom du programme de test", "Test program name"),
            value=default_test_name,
            help=T("Commence par 'M', max 8 caractères", "Starts with 'M', max 8 chars")
        ).upper()
        
        st.markdown('</div>', unsafe_allow_html=True)

        def valid_prog_name(name: str) -> Tuple[bool, str]:
            if not name:
                return False, T("Nom vide", "Empty name")
            if len(name) > 8:
                return False, T("Plus de 8 caractères", "More than 8 chars")
            if name[0] != "B":
                return False, T("Doit commencer par 'B'", "Must start with 'B'")
            if not all(ch.isalnum() for ch in name):
                return False, T("Caractères invalides", "Invalid characters")
            return True, ""

        ok_name, err_name = valid_prog_name(test_prog_name)
        if not ok_name:
            st.markdown(f"""
            <div class="warning-box">
                ℹ️ Nom invalide : {err_name}
            </div>
            """, unsafe_allow_html=True)

        st.markdown(f"""
        <div class="info-box">
            🔍 Type : <strong>{language_type}</strong> | Programme test : <strong>{test_prog_name}</strong>
        </div>
        """, unsafe_allow_html=True)

        gen_disabled = not module_starts_ok or not ok_name
        
        if st.button(
            "🚀 " + T("GÉNÉRER TEST + SCÉNARIOS", "GENERATE TEST + SCENARIOS"),
            disabled=gen_disabled,
            use_container_width=True
        ):
            prompt_text = get_prompt(
                "TEST_FACTORY", language_type,
                source_code=source_code,
                module_base=module_base,
                test_prog_name=test_prog_name
            )
            
            client = llm_client(max_tokens=4000, temperature=0.2)
            if client:
                with st.spinner(T("🧠 Génération en cours...", "🧠 Generating...")):
                    try:
                        response = client.invoke(prompt_text)
                        result = response.content if hasattr(response, 'content') else str(response)
                        
                        if result:
                            # Parse sections
                            cobol_test, cases_raw = "", ""
                            if "=== COBOL_TEST ===" in result:
                                after = result.split("=== COBOL_TEST ===", 1)[1]
                                if "=== CAS_DE_TEST ===" in after:
                                    cobol_test, cases_raw = after.split("=== CAS_DE_TEST ===", 1)
                                else:
                                    cobol_test = after

                            cobol_test_clean = cobol_test.strip()
                            
                            # Header injection
                            AUTHOR_NAME = "AYMANE"
                            DATE_STR = datetime.now().strftime("%Y/%m/%d")
                            header = f"""      *> PROGRAM-ID : {test_prog_name}
      *> AUTHOR     : {AUTHOR_NAME}
      *> DATE       : {DATE_STR}
"""
                            cobol_with_header = header + cobol_test_clean if cobol_test_clean else ""

                            # Display COBOL
                            if cobol_with_header:
                                st.markdown('<div class="glass-card">', unsafe_allow_html=True)
                                st.subheader("📘 " + T("Programme COBOL de test", "COBOL Test Program"))
                                st.code(cobol_with_header, language="cobol")
                                st.markdown('</div>', unsafe_allow_html=True)
                                
                                st.download_button(
                                    "💾 " + T("Télécharger COBOL (.cbl)", "Download COBOL (.cbl)"),
                                    data=cobol_with_header.encode("utf-8"),
                                    file_name=f"{test_prog_name}.cbl",
                                    use_container_width=True
                                )

                            # Parse scenarios
                            rows, headers = [], []
                            for ln in cases_raw.splitlines():
                                line = ln.strip()
                                if not line or line.startswith("#"):
                                    continue
                                parts = [c.strip() for c in line.split("|")]
                                if not headers and "Nom_Scenario" in line:
                                    headers = parts
                                    continue
                                if headers and len(parts) == len(headers):
                                    rows.append(parts)

                            if headers and rows:
                                df_cases = pd.DataFrame(rows, columns=headers)
                            else:
                                df_cases = pd.DataFrame(columns=[
                                    "Nom_Scenario", "Condition_Entree", "Donnees_Exemple",
                                    "Resultat_Attendu", "RC_Attendu", "Artefacts_Impactes", "Commentaires"
                                ])

                            # Display matrix
                            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
                            st.subheader("🧾 " + T("Matrice de scénarios", "Scenario Matrix"))
                            
                            if not df_cases.empty:
                                st.dataframe(df_cases, use_container_width=True, height=400)
                            else:
                                st.warning(T("Aucun scénario détecté", "No scenarios detected"))
                            
                            st.markdown('</div>', unsafe_allow_html=True)

                            # Export Excel
                            scenarios_name = f"SCENARIOS_{module_base}"
                            excel_buf = io.BytesIO()
                            
                            try:
                                with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
                                    df_cases.to_excel(writer, index=False, sheet_name="SCENARIOS")
                                excel_buf.seek(0)
                                
                                st.download_button(
                                    "📥 " + T("Télécharger Scénarios (Excel)", "Download Scenarios (Excel)"),
                                    data=excel_buf,
                                    file_name=f"{scenarios_name}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True
                                )
                            except ImportError:
                                csv_buf = io.BytesIO()
                                df_cases.to_csv(csv_buf, index=False, encoding="utf-8")
                                csv_buf.seek(0)
                                
                                st.download_button(
                                    "📥 " + T("Télécharger Scénarios (CSV)", "Download Scenarios (CSV)"),
                                    data=csv_buf,
                                    file_name=f"{scenarios_name}.csv",
                                    use_container_width=True
                                )
                                
                    except Exception as e:
                        st.error(f"❌ Erreur: {e}")

# ===================== FOOTER PRO =====================
st.markdown("""
<div class="footer-pro">
    <div class="footer-title">💼 Modernisation Mainframe IBM A&T</div>
    <div class="footer-team">
        👥 <span>Équipe</span> : Youness • Hanane • Nezha • Aimane • Khaoula • Naoufal • Imane • Mariem
    </div>
    <div style="margin-top: 1rem; color: #666; font-size: 0.85rem;">
         © 2025 Tous droits réservés
    </div>
</div>
""", unsafe_allow_html=True)